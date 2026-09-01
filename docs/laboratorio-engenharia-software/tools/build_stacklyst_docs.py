from __future__ import annotations

import hashlib
import json
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


REPO = Path(r"C:\Users\PEDRO\Documents\DevDeck")
ROOT = REPO / "docs" / "laboratorio-engenharia-software"
OUTPUT = ROOT / "entregaveis"
DIAGRAMS = ROOT / "diagramas"
SCREENSHOTS = REPO / "public" / "screenshots"
CURRENT_PROTOTYPES = ROOT / "prototipos" / "atuais"
WORK = Path(r"C:\Users\PEDRO\AppData\Local\Temp\codex-stacklyst-docs-20260821")

TEMPLATES = {
    "vision": WORK / "01-visao-template.docx",
    "activities": WORK / "02-atividades-template.docx",
    "requirements": WORK / "03-requisitos-template.docx",
    "use_cases": WORK / "04-casos-template.docx",
}

FINAL_FILES = {
    "vision": OUTPUT / "01-documento-de-visao-stacklyst.docx",
    "activities": OUTPUT / "02-atividades-do-negocio-stacklyst.docx",
    "requirements": OUTPUT / "03-requisitos-do-sistema-stacklyst.docx",
    "use_cases": OUTPUT / "04-casos-de-uso-stacklyst.docx",
}

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PURPLE = "6D28D9"
LIGHT_PURPLE = "EDE9FE"
GREEN = "166534"
LIGHT_GREEN = "DCFCE7"
AMBER = "92400E"
LIGHT_AMBER = "FEF3C7"
RED = "991B1B"
LIGHT_RED = "FEE2E2"
GRAY = "E5E7EB"
WHITE = "FFFFFF"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width: float) -> None:
    """Fix the table width in twips so Word does not recalculate its grid."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Atualize o campo no Word"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    sizes = {1: 18, 2: 14, 3: 12}
    for level, size in sizes.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("111827")
        style.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
        # The legacy templates attach automatic numbering to Heading styles.
        # Headings in this generator already carry the official numbering, so
        # removing numPr prevents results such as "3. 3. Objetivo".
        p_pr = style._element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_header_footer(document: Document, document_title: str) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False
    # Os templates oficiais possuem textos de exemplo no cabeçalho e no rodapé.
    # Removê-los evita duplicar placeholders ou atribuir autoria não confirmada.
    for child in list(header._element):
        header._element.remove(child)
    table = header.add_table(rows=1, cols=2, width=Inches(6.85))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 6.85)
    set_cell_width(table.cell(0, 0), 3.4)
    set_cell_width(table.cell(0, 1), 3.45)
    left = table.cell(0, 0).paragraphs[0]
    left.add_run("Stacklyst").bold = True
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(document_title)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)

    footer = section.footer
    footer.is_linked_to_previous = False
    for child in list(footer._element):
        footer._element.remove(child)
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.85))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    set_table_width(footer_table, 6.85)
    set_cell_width(footer_table.cell(0, 0), 3.6)
    set_cell_width(footer_table.cell(0, 1), 3.25)
    p_left = footer_table.cell(0, 0).paragraphs[0]
    p_left.add_run("Versão do documento: 1.0")
    p_right = footer_table.cell(0, 1).paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run("Página ")
    add_field(p_right, "PAGE")
    p_right.add_run(" de ")
    add_field(p_right, "NUMPAGES")
    for cell in footer_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)


def add_cover(document: Document, subtitle: str) -> None:
    for _ in range(7):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Stacklyst")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title.paragraph_format.space_after = Pt(14)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.bold = True
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    meta.add_run("Laboratório de Engenharia de Software\n").bold = True
    meta.add_run("Projeto acadêmico — [ANOTAÇÃO — VALIDAR COM A PROFESSORA]\n")
    meta.add_run("Data: [ANOTAÇÃO — DATA A DEFINIR]")
    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("1. Índice", level=1)
    p = document.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    document.add_page_break()


def add_note(document: Document, label: str, text: str, level: str = "yellow", spacing: bool = True) -> None:
    colors = {
        "red": (LIGHT_RED, RED),
        "yellow": (LIGHT_AMBER, AMBER),
        "blue": (LIGHT_BLUE, BLUE),
    }
    fill, color = colors[level]
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 6.75)
    set_row_cant_split(table.rows[0])
    set_cell_width(table.cell(0, 0), 6.75)
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    if spacing:
        document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(document: Document, items: Iterable[str], level: int = 0) -> None:
    for item in items:
        p = document.add_paragraph()
        left = 0.25 + (0.25 * level)
        p.paragraph_format.left_indent = Inches(left)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run("• ")
        p.add_run(item)


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for index, item in enumerate(items, 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(f"{index}. ").bold = True
        p.add_run(item)


def add_labeled(document: Document, label: str, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.keep_together = True
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 8.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, sum(widths) if widths else 6.75)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, BLUE)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(font_size)
        if widths:
            set_cell_width(cell, widths[idx])
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                set_cell_width(cells[idx], widths[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document: Document, image_path: Path, caption: str, max_width=6.65, max_height=8.2) -> None:
    if not image_path.exists():
        add_note(document, "[ANOTAÇÃO — ARQUIVO NÃO ENCONTRADO]", str(image_path), "red")
        return
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    scale = min(max_width / width_px, max_height / height_px)
    width = width_px * scale
    height = height_px * scale
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)


def add_code_block(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 6.7)
    set_cell_width(table.cell(0, 0), 6.7)
    set_cell_shading(table.cell(0, 0), "F3F4F6")
    cell = table.cell(0, 0)
    cell.text = ""
    for line in code.rstrip().splitlines():
        p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(6.5)


def new_from_template(key: str, title: str, subtitle: str) -> Document:
    source = TEMPLATES[key]
    target = FINAL_FILES[key]
    if not source.exists():
        raise FileNotFoundError(f"Template convertido não encontrado: {source}")
    OUTPUT.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)
    document = Document(target)
    clear_body(document)
    configure_styles(document)
    set_header_footer(document, title)
    add_cover(document, subtitle)
    add_toc(document)
    return document


SCOPE_MODULES = [
    ("Autenticação e usuários", "Cadastro, login, logout, recuperação de acesso, sessão e controle de papéis.", "Implementado", "src/app/api/auth; src/lib/auth.ts; src/lib/supabase"),
    ("Perfil", "Dados públicos do desenvolvedor, histórico de publicações e indicadores de evolução.", "Implementado", "src/app/profile/[username]; src/app/api/profile"),
    ("Trilhas", "Percursos por linguagem, preferências de curso, níveis e checkpoints.", "Implementado", "src/app/trails; src/lib/trails; src/app/api/trails"),
    ("Lições", "Páginas próprias com conteúdo e etapas interativas.", "Implementado", "src/app/lesson/[lessonId]; src/lib/lessons"),
    ("Exercícios", "Múltipla escolha, lacunas, ordenação, associação, terminal, editor e blocos.", "Implementado", "src/lib/lessons/evaluators.ts"),
    ("Desafios e execução de código", "Problemas de programação, envio de solução e execução em serviço externo.", "Implementado", "src/app/api/run; src/lib/code-runner.ts"),
    ("Duelos", "Convites, aceitação, desafio temporizado, soluções, votos e resultado.", "Implementado parcialmente", "src/app/duels; src/services/duel.service.ts"),
    ("Matchmaking", "Busca imediata por faixa de XP com expansão progressiva; não há fila persistente de espera.", "Implementado parcialmente", "src/services/duel.service.ts"),
    ("Avaliações", "Avaliação automática e fluxo de avaliação humana para duelos.", "Implementado parcialmente", "src/app/api/evaluations; src/services/evaluator.service.ts"),
    ("Avaliadores", "Elegibilidade, candidatura, aprovação administrativa e parecer técnico.", "Implementado", "src/app/evaluators; src/app/evaluations; src/services/evaluator.service.ts"),
    ("XP, níveis e sequência", "Pontuação transacional, níveis e ofensiva por atividade.", "Implementado", "src/services/xp.service.ts; src/lib/streak.ts"),
    ("Conquistas", "Badges por critérios de atividade e sequência.", "Implementado", "prisma/schema.prisma; src/services/xp.service.ts"),
    ("Ranking e divisões", "Ranking global/por linguagem e faixas Bronze a Diamante usadas nos duelos.", "Implementado parcialmente", "src/app/api/leaderboard; src/services/duel.service.ts"),
    ("Comunidade", "Feed, publicações, respostas, reações, votos, favoritos, seguidores e mensagens.", "Implementado", "src/app/feed; src/services/post.service.ts; src/app/api/posts"),
    ("Eventos", "Listagem, criação e participação em eventos/desafios comunitários.", "Implementado parcialmente", "src/app/events; src/services/event.service.ts"),
    ("Notificações", "Central web e tipos de eventos da plataforma.", "Implementado na web", "src/app/notifications; src/services/notification.service.ts"),
    ("Aplicação mobile", "Recebimento de push e consulta de informações selecionadas.", "Planejado", "[ANOTAÇÃO — DECISÃO TÉCNICA PENDENTE]"),
    ("Vagas e empresas", "Cadastro de empresa, publicação, consulta e candidatura a oportunidades.", "Implementado parcialmente", "src/app/jobs; src/app/recruiter; src/services/job.service.ts"),
    ("Administração", "Usuários, papéis, denúncias, avaliadores, métricas e geração administrativa de quiz.", "Implementado parcialmente", "src/app/admin; src/app/api/admin"),
    ("Acompanhamento de progresso", "Tentativas, XP, níveis, ofensiva, checkpoints e atividade semanal.", "Implementado", "QuizAttempt, LanguageTrail; src/app/api/quiz/weekly-activity"),
]


RISKS = [
    ("Atraso de desenvolvimento", "Gerencial", "Alta", "Alto", "Alto", "Planejar entregas incrementais e priorizar o MVP.", "Reduzir itens de prioridade média/baixa e replanejar marcos."),
    ("Mudança frequente de requisitos", "Requisitos", "Alta", "Alto", "Alto", "Manter baseline versionado e registrar decisões.", "Executar análise de impacto e nova validação acadêmica."),
    ("Escopo excessivo", "Gerencial", "Alta", "Alto", "Alto", "Separar MVP, expansão e itens fora do escopo.", "Congelar novas funcionalidades até concluir o núcleo."),
    ("Dependência excessiva de IA", "Técnico", "Média", "Alto", "Alto", "Manter conteúdo curado, fallbacks e decisão humana.", "Desativar o provedor sem interromper fluxos essenciais."),
    ("Resposta incorreta ou inadequada da IA", "Qualidade", "Alta", "Alto", "Alto", "Validar formato, exibir caráter auxiliar e limitar contexto.", "Permitir denúncia, revisar prompts e remover resposta inadequada."),
    ("Custo de infraestrutura", "Financeiro", "Média", "Alto", "Alto", "Monitorar uso e definir cotas por serviço.", "Reduzir recursos opcionais e migrar para plano compatível."),
    ("Falha de segurança", "Segurança", "Média", "Alto", "Alto", "RBAC, validação, rate limit, segredo server-only e testes.", "Revogar credenciais, conter acesso e corrigir vulnerabilidade."),
    ("Exposição indevida de dados", "Privacidade", "Média", "Alto", "Alto", "Definir campos públicos e consentimento para recrutamento.", "Restringir consultas, registrar incidente e revisar permissões."),
    ("Baixa adesão", "Negócio", "Média", "Alto", "Alto", "Testar usabilidade e valor das trilhas com usuários.", "Simplificar onboarding e ajustar mecânicas com feedback."),
    ("Complexidade do matchmaking", "Técnico", "Alta", "Médio", "Alto", "Começar com regras determinísticas e métricas.", "Usar desafio direto e busca ampliada como alternativa."),
    ("Falha na avaliação automática de código", "Técnico", "Média", "Alto", "Alto", "Sandbox externo, timeout, casos de teste e fallback humano.", "Suspender pontuação e encaminhar para revisão."),
    ("Problemas de desempenho", "Técnico", "Média", "Alto", "Alto", "Paginação, índices, cache e testes de carga.", "Reduzir consultas, limitar listas e escalar serviços críticos."),
    ("Indisponibilidade dos integrantes", "Organizacional", "Média", "Alto", "Alto", "Distribuir conhecimento e registrar decisões.", "Repriorizar entregas e redistribuir responsabilidades."),
    ("Integração web/mobile", "Técnico", "Média", "Médio", "Médio", "Definir contrato de API e estratégia de notificações.", "Manter notificações web e adiar push mobile."),
    ("Moderação insuficiente", "Operacional", "Média", "Alto", "Alto", "Denúncia, fila, critérios e auditoria administrativa.", "Ocultar conteúdo, restringir usuário e ampliar revisão."),
    ("Fraude em ranking ou duelo", "Segurança", "Alta", "Alto", "Alto", "Idempotência, validação server-side e registros de auditoria.", "Anular resultado, recalcular pontuação e bloquear abuso."),
]


MILESTONES = [
    "Início do projeto", "Levantamento de requisitos", "Documento de visão", "Atividades do negócio",
    "Requisitos do sistema", "Casos de uso", "Prototipação", "Modelagem do banco", "Definição de arquitetura",
    "MVP", "Módulo de autenticação", "Trilhas e lições", "Gamificação", "Duelos", "IA auxiliar",
    "Comunidade", "Vagas e empresas", "Testes", "Testes de aceitação", "Correções", "Deploy", "Entrega final",
]


def build_vision() -> None:
    document = new_from_template("vision", "Documento de Visão", "Documento de Visão")
    document.add_heading("2. Objetivo", level=1)
    document.add_paragraph(
        "O Stacklyst é uma plataforma acadêmica gamificada para ensino, prática e evolução em programação. "
        "O produto reúne trilhas, lições, exercícios, desafios, duelos, comunidade e acompanhamento de progresso "
        "em um ambiente integrado destinado principalmente a estudantes e desenvolvedores em formação ou aperfeiçoamento."
    )
    document.add_paragraph(
        "O problema central é a fragmentação entre aprender conceitos, praticar código, manter constância e demonstrar evolução. "
        "O objetivo educacional é favorecer prática ativa, feedback frequente e progressão observável; o objetivo da plataforma "
        "é conectar essas experiências a interações comunitárias e, quando autorizado, a oportunidades profissionais."
    )
    document.add_paragraph(
        "A gamificação utiliza XP, níveis, sequência de atividades, conquistas, rankings e divisões como mecanismos de engajamento, "
        "sem substituir a aprendizagem. A inteligência artificial atua apenas como apoio para explicar dúvidas, sugerir conteúdo ou "
        "gerar materiais auxiliares. Ela não é autoridade absoluta em decisões que exijam avaliação humana, moderação, autorização ou impacto competitivo."
    )

    document.add_heading("3. Necessidade do Negócio", level=1)
    document.add_paragraph(
        "O projeto não possui empresa cliente real informada. Para fins acadêmicos, o Stacklyst é caracterizado como uma plataforma educacional "
        "voltada à aprendizagem prática de programação. [ANOTAÇÃO — VALIDAR COM A PROFESSORA]: confirmar se o template exige a identificação de um cliente fictício adicional."
    )
    document.add_paragraph(
        "O cenário observado contém plataformas muito passivas, exercícios pouco variados, dificuldade de manter frequência nos estudos, pouca interação entre estudantes, "
        "distância entre conteúdo e situações práticas, dificuldade para medir a evolução e pouca conexão entre aprendizagem, comunidade e oportunidades profissionais."
    )
    document.add_paragraph(
        "O Stacklyst pretende melhorar esse cenário com percursos organizados, atividades interativas, execução de código, feedback, desafios competitivos, avaliação humana quando necessária, "
        "progressão gamificada, espaços comunitários e um módulo de oportunidades. A solução mantém uma separação entre evidência educacional e decisão de recrutamento: desempenho público pode ser um indicador, mas não uma decisão automática."
    )

    document.add_heading("4. Descrição do Escopo", level=1)
    document.add_heading("4.1 Dentro do escopo", level=2)
    add_table(document, ["Módulo", "Descrição", "Estado", "Evidência"], [list(row) for row in SCOPE_MODULES], [1.35, 2.75, 1.05, 1.6], 7.2)
    document.add_heading("4.2 Fora do escopo da primeira versão", level=2)
    add_bullets(document, [
        "Marketplace de cursos, serviços ou produtos digitais.",
        "Contratação, assinatura de contrato e pagamento dentro do Stacklyst.",
        "IDE online completa equivalente a um ambiente desktop, com depurador e gerenciamento avançado de projetos.",
        "Sistema próprio de videoconferência.",
        "Emissão de certificações oficiais reconhecidas por órgãos externos.",
        "Integração imediata com dezenas de plataformas externas.",
        "Decisões automáticas de contratação, reprovação, moderação ou punição baseadas somente em IA.",
    ])
    document.add_paragraph("Esses itens poderão ser considerados em expansões futuras mediante nova análise de requisitos, riscos, custos e aprovação acadêmica.")

    document.add_heading("5. Equipe", level=1)
    add_note(document, "[ANOTAÇÃO — PREENCHER INTEGRANTES]", "Duplicar a ficha abaixo para cada integrante real; não foram fornecidos nomes, formações ou experiências.", "yellow")
    add_table(document, ["Campo", "Preenchimento"], [
        ["Nome", "[ANOTAÇÃO — PREENCHER INTEGRANTE]"],
        ["Formação", "[ANOTAÇÃO — PREENCHER INTEGRANTE]"],
        ["Experiência", "[ANOTAÇÃO — PREENCHER INTEGRANTE]"],
        ["Papel no projeto", "[ANOTAÇÃO — PREENCHER INTEGRANTE]"],
        ["Responsabilidades", "[ANOTAÇÃO — PREENCHER INTEGRANTE]"],
    ], [1.7, 5.05], 9)

    document.add_heading("6. Especificações Técnicas", level=1)
    document.add_paragraph("As tecnologias abaixo foram verificadas no repositório. Itens não presentes no código permanecem como decisão pendente ou sugestão.")
    tech_rows = [
        ["Frontend Web", "Next.js 16.3.2 (App Router), React 19.2.4, TypeScript 5, Tailwind CSS 4, CodeMirror e Tiptap.", "package.json; src/app"],
        ["Backend", "Route Handlers do Next.js em runtime Node, serviços TypeScript, validação Zod e controles de taxa.", "src/app/api; src/services; src/lib/api-handler.ts"],
        ["Banco de Dados", "PostgreSQL hospedável no Supabase, modelado pelo Prisma ORM 7.9.1.", "prisma/schema.prisma; docs/DATABASE.md"],
        ["Autenticação", "Supabase Auth/SSR com sessão verificada e camada server-only de JWT da aplicação.", "src/lib/supabase; src/lib/auth.ts; src/lib/jwt.ts"],
        ["Aplicação Mobile", "[ANOTAÇÃO — DECISÃO TÉCNICA PENDENTE]: não existe aplicativo mobile no repositório auditado.", "Nenhuma implementação mobile localizada"],
        ["Infraestrutura", "Docker/Docker Compose; documentação de Vercel e Supabase; Upstash Redis opcional para rate limit distribuído.", "Dockerfile; docker-compose*.yml; docs/DEPLOYMENT.md"],
        ["APIs", "REST/JSON internas; Supabase; execução de código por Wandbox/Judge0.", "src/app/api; src/lib/code-runner.ts"],
        ["Controle de versão", "Git, GitHub, Conventional Commits, Husky, lint-staged e CI configurável.", ".git; commitlint.config.js; .husky; package.json"],
        ["Arquitetura geral", "Aplicação web modular: UI e Server Components → API Routes → serviços/domínio → Prisma/PostgreSQL; integrações externas isoladas em bibliotecas.", "docs/ARCHITECTURE.md; estrutura src"],
    ]
    add_table(document, ["Camada", "Especificação verificada", "Evidência"], tech_rows, [1.35, 3.8, 1.6], 7.7)
    add_note(document, "[SUGESTÃO TÉCNICA]", "Antes de iniciar o mobile, definir contrato de API, autenticação, estratégia de push, armazenamento seguro e compatibilidade com as permissões de privacidade.", "blue")

    document.add_heading("7. Riscos", level=1)
    add_table(document, ["Risco", "Categoria", "Prob.", "Impacto", "Nível", "Estratégia", "Contingência"], [list(r) for r in RISKS], [1.2, 0.75, 0.5, 0.55, 0.5, 1.65, 1.65], 6.4)

    document.add_heading("8. Cronograma de Marcos Resumido", level=1)
    document.add_paragraph("As datas oficiais não foram informadas. Os marcos abaixo preservam a ordem acadêmica e de desenvolvimento sugerida.")
    add_table(document, ["Marco", "Data"], [[m, "[ANOTAÇÃO — DATA A DEFINIR]"] for m in MILESTONES], [4.6, 2.15], 8.5)

    document.add_heading("9. Orçamento Resumido", level=1)
    budget = [
        ("Custos fixos", "Hardware", "R$ [VALOR A DEFINIR]", "Verificar equipamentos já disponíveis."),
        ("Custos fixos", "Licenças de software", "R$ [VALOR A DEFINIR]", "Ferramentas gratuitas podem ser consideradas, sem confirmação de elegibilidade."),
        ("Custos fixos", "Treinamentos", "R$ [VALOR A DEFINIR]", "[ANOTAÇÃO — ORÇAMENTO A VALIDAR]"),
        ("Infraestrutura", "Hospedagem e domínio", "R$ [VALOR A DEFINIR]", "Planos gratuitos podem ser avaliados, sujeitos a limites."),
        ("Infraestrutura", "Banco de dados e armazenamento", "R$ [VALOR A DEFINIR]", "Dimensionar após teste de carga."),
        ("Infraestrutura", "APIs e execução de código", "R$ [VALOR A DEFINIR]", "Considerar cotas e volume."),
        ("Infraestrutura", "Serviços de IA", "R$ [VALOR A DEFINIR]", "Comparar inferência local e provedores externos."),
        ("Custos variáveis", "Desenvolvimento", "R$ [VALOR A DEFINIR]", "Não foram informados salários ou horas."),
        ("Custos variáveis", "Pessoal de apoio", "R$ [VALOR A DEFINIR]", "[ANOTAÇÃO — ORÇAMENTO A VALIDAR]"),
        ("Custos variáveis", "Encargos e benefícios", "R$ [VALOR A DEFINIR]", "[ANOTAÇÃO — ORÇAMENTO A VALIDAR]"),
        ("Custos variáveis", "Instalações, energia, materiais e rede", "R$ [VALOR A DEFINIR]", "[ANOTAÇÃO — ORÇAMENTO A VALIDAR]"),
        ("Qualidade", "Testes e homologação", "R$ [VALOR A DEFINIR]", "Incluir ambientes e dispositivos."),
        ("Riscos", "Margem de contingência", "R$ [VALOR A DEFINIR]", "Percentual ainda não definido."),
        ("Financeiro", "Margem de lucro", "R$ [VALOR A DEFINIR]", "Validar aplicabilidade em projeto acadêmico."),
    ]
    add_note(document, "[ANOTAÇÃO — ORÇAMENTO A VALIDAR]", "Nenhum valor financeiro foi fornecido; todos os totais dependem de levantamento real e orientação acadêmica.", "yellow")
    add_table(document, ["Categoria", "Item", "Valor", "Observação"], [list(r) for r in budget], [1.1, 2.35, 1.45, 1.85], 7.0)

    document.save(FINAL_FILES["vision"])


ACTIVITIES = [
    {
        "id": "AN01", "name": "Realizar aprendizado por trilha",
        "objective": "Permitir que o usuário escolha um percurso, conclua lições e exercícios e acompanhe a evolução.",
        "participants": "Usuário / Estudante / Desenvolvedor; Stacklyst; Executor de código quando necessário.",
        "preconditions": "Usuário autenticado; trilha publicada; conteúdo e pré-requisitos disponíveis.",
        "flow": ["O usuário seleciona uma trilha.", "O Stacklyst exibe a estrutura, o progresso e abre a lição escolhida ou seguinte.", "O usuário estuda a lição e resolve o exercício associado.", "O Stacklyst avalia a resposta; quando incorreta, fornece feedback e permite nova tentativa.", "Quando correta, o sistema registra a tentativa, atualiza progresso e verifica se é a última lição.", "Se ainda houver conteúdo, a próxima lição é liberada e o usuário continua; caso contrário, a conclusão da trilha é registrada."],
        "decisions": "Resposta correta? É a última lição? Os pré-requisitos da próxima lição foram cumpridos?",
        "exceptions": "Conteúdo indisponível, falha no executor de código, perda de conexão ou tentativa já pontuada. O sistema deve preservar o progresso confirmado e informar o problema.",
        "result": "Progresso, tentativa, XP elegível e conclusão ficam registrados; a próxima etapa é apresentada.",
        "evidence": "src/app/trails; src/app/lesson/[lessonId]; src/lib/lessons; src/app/api/trails/checkpoint",
        "diagram": "AN01-aprendizado-por-trilha",
    },
    {
        "id": "AN02", "name": "Participar de duelo",
        "objective": "Organizar a procura de oponente, o convite, a resolução do desafio, a avaliação e o resultado competitivo.",
        "participants": "Desafiante; Oponente; Stacklyst; Avaliador quando exigido.",
        "preconditions": "Usuários autenticados e sem bloqueio/cooldown; linguagem selecionada; desafio disponível.",
        "flow": ["O desafiante solicita duelo direto ou busca automática.", "O Stacklyst verifica se o usuário pode duelar e busca um oponente compatível.", "Se nenhum oponente for localizado, informa indisponibilidade; se houver, envia um convite.", "O oponente analisa e aceita ou rejeita. A rejeição é registrada e pode gerar penalização conforme regra validada.", "Com o aceite, o sistema cria o desafio e ambos enviam soluções dentro do prazo aplicável.", "O sistema recebe as soluções e decide se a avaliação automática é suficiente ou se deve encaminhar a um avaliador.", "A avaliação é registrada, o duelo é fechado e os participantes consultam o resultado."],
        "decisions": "Pode duelar? Há oponente? O convite foi aceito? O tempo expirou? Avaliação humana é necessária?",
        "exceptions": "Nenhum oponente, convite expirado, rejeição, cooldown, solução não enviada, executor indisponível ou conflito de avaliação.",
        "result": "Duelo fechado ou registrado como não concluído; resultado, notificações e efeitos competitivos são processados conforme regras aprovadas.",
        "evidence": "src/services/duel.service.ts; src/app/api/duels; src/app/api/evaluations",
        "diagram": "AN02-participar-de-duelo",
    },
    {
        "id": "AN03", "name": "Resolver exercício",
        "objective": "Avaliar a prática do usuário em diferentes formatos e oferecer feedback verificável.",
        "participants": "Usuário; Stacklyst; Executor de código quando o tipo exigir.",
        "preconditions": "Exercício acessível e dados da sessão carregados.",
        "flow": ["O usuário escolhe um exercício e o sistema carrega o tipo de interação.", "O usuário informa ou monta a solução.", "O sistema valida o preenchimento e, quando houver código, solicita execução ao serviço configurado.", "O sistema compara resposta, saída ou testes com o resultado esperado.", "Se incorreta, exibe feedback e o usuário revisa a solução.", "Se correta, registra a tentativa, concede a pontuação elegível e atualiza progresso.", "O usuário consulta o resultado."],
        "decisions": "Entrada válida? Há código para executar? Serviço disponível? Solução correta? A tentativa já recebeu XP?",
        "exceptions": "Entrada incompleta, sintaxe inválida, timeout, executor indisponível, resposta duplicada ou falha de persistência.",
        "result": "Tentativa e feedback registrados; XP e progresso atualizados apenas quando a regra permitir.",
        "evidence": "src/lib/lessons/evaluators.ts; src/lib/code-runner.ts; src/app/api/quiz/[id]/attempt",
        "diagram": "AN03-resolver-exercicio",
    },
    {
        "id": "AN04", "name": "Evoluir na gamificação",
        "objective": "Converter atividades válidas em progressão observável sem duplicar recompensas.",
        "participants": "Usuário; Stacklyst.",
        "preconditions": "Atividade concluída e usuário identificado.",
        "flow": ["O usuário conclui uma atividade.", "O sistema verifica elegibilidade e se a atividade já foi pontuada.", "Quando elegível e inédita, concede XP de forma transacional.", "O sistema recalcula nível e sequência de atividade.", "Verifica critérios de conquistas e registra as que foram alcançadas.", "Atualiza ranking/divisão aplicável.", "O usuário consulta a evolução e recebe eventuais conquistas."],
        "decisions": "Atividade elegível? Já pontuada? Houve mudança de nível? Critério de conquista atingido?",
        "exceptions": "Atividade não elegível, tentativa duplicada ou falha transacional; a pontuação anterior deve ser mantida.",
        "result": "XP, nível, sequência, conquistas e ranking permanecem consistentes.",
        "evidence": "src/services/xp.service.ts; src/lib/streak.ts; src/app/api/leaderboard",
        "diagram": "AN04-evoluir-na-gamificacao",
    },
    {
        "id": "AN05", "name": "Interagir com a comunidade",
        "objective": "Permitir publicação, colaboração e moderação de conteúdo técnico.",
        "participants": "Autor; Comunidade; Stacklyst; Administrador quando houver denúncia.",
        "preconditions": "Usuário autenticado para publicar/interagir; regras de conteúdo disponíveis.",
        "flow": ["O autor redige e envia uma publicação.", "O Stacklyst valida o conteúdo; se houver erro, solicita correção.", "O conteúdo válido é publicado e exibido no feed.", "A comunidade visualiza, reage, vota, responde ou denuncia.", "Interações válidas são registradas e geram notificações.", "Quando denunciado, o conteúdo é encaminhado ao administrador.", "O administrador analisa a denúncia e aplica a decisão cabível."],
        "decisions": "Conteúdo válido? A interação é permitida? Houve denúncia? A denúncia procede?",
        "exceptions": "Conteúdo vazio/inválido, limite de publicação, usuário sem permissão, duplicidade ou indisponibilidade.",
        "result": "Publicação e interações registradas; denúncia tratada quando aplicável.",
        "evidence": "src/app/feed; src/app/api/posts; src/app/api/admin/reports",
        "diagram": "AN05-interagir-com-comunidade",
    },
    {
        "id": "AN07", "name": "Atuar como avaliador",
        "objective": "Credenciar usuários elegíveis e registrar avaliações humanas auditáveis.",
        "participants": "Candidato; Administrador; Avaliador; Stacklyst.",
        "preconditions": "Candidato autenticado; critérios de elegibilidade verificáveis; administrador disponível para decisão.",
        "flow": ["O candidato solicita a função de avaliador.", "O Stacklyst verifica critérios; se não atendidos, informa inelegibilidade.", "O candidato elegível informa motivação e tecnologias e a candidatura é registrada.", "O administrador analisa e aprova ou rejeita, registrando a decisão.", "Quando aprovado, o sistema atribui a função e disponibiliza avaliações pendentes.", "O avaliador seleciona uma pendência, revisa as soluções e emite parecer.", "O sistema registra a avaliação, fecha o desafio e notifica os envolvidos."],
        "decisions": "Critério atendido? Há candidatura pendente? Candidatura aprovada? Avaliação completa e autorizada?",
        "exceptions": "Usuário inelegível, candidatura duplicada, administrador sem permissão, duelo inexistente ou dados de avaliação inválidos.",
        "result": "Candidatura decidida e, quando aplicável, avaliação humana registrada com rastreabilidade.",
        "evidence": "src/services/evaluator.service.ts; src/app/api/evaluators; src/app/api/admin/evaluators; src/app/api/evaluations",
        "diagram": "AN07-atuar-como-avaliador",
    },
    {
        "id": "AN08", "name": "Empresa divulgar vaga",
        "objective": "Permitir que empresas autorizadas publiquem oportunidades e acompanhem candidaturas.",
        "participants": "Empresa / Recrutador; Usuário; Stacklyst.",
        "preconditions": "Conta autenticada com papel autorizado; empresa cadastrada e vinculada ao responsável.",
        "flow": ["A empresa ou recrutador cadastra a organização.", "O Stacklyst verifica a autorização.", "O recrutador redige a oportunidade e solicita publicação.", "O sistema valida a vaga; se inválida, solicita correção; se válida, disponibiliza a oportunidade.", "O usuário consulta, analisa e envia candidatura.", "O sistema registra a candidatura e exibe somente dados permitidos.", "A empresa acompanha e atualiza o processo; o candidato é notificado e consulta a situação."],
        "decisions": "Cadastro autorizado? Vaga válida e aberta? Usuário já se candidatou? Dados podem ser exibidos à empresa?",
        "exceptions": "Empresa não autorizada, vaga incompleta/encerrada, candidatura duplicada ou acesso a dado não permitido.",
        "result": "Vaga publicada e candidatura acompanhável, respeitando permissão e privacidade.",
        "evidence": "src/app/api/companies; src/app/api/jobs; src/services/job.service.ts; src/app/recruiter",
        "diagram": "AN08-empresa-divulgar-vaga",
    },
    {
        "id": "AN09", "name": "Administração da plataforma",
        "objective": "Executar ações administrativas autorizadas com validação, auditoria e comunicação aos envolvidos.",
        "participants": "Administrador; Stacklyst.",
        "preconditions": "Administrador autenticado e com papel válido.",
        "flow": ["O administrador acessa o painel.", "O Stacklyst valida a permissão e carrega informações.", "O administrador seleciona usuários, permissões, denúncias, avaliadores, eventos ou configurações.", "O administrador analisa os dados, define e confirma a ação.", "O sistema valida a consistência da alteração.", "Se inválida, informa a inconsistência; se válida, executa a operação.", "O Stacklyst registra auditoria e notifica os envolvidos quando necessário."],
        "decisions": "Permissão válida? Área suportada? Ação consistente? Exige confirmação ou justificativa?",
        "exceptions": "Acesso negado, dado desatualizado, ação inválida, operação concorrente ou indisponibilidade.",
        "result": "Ação administrativa registrada e comunicada, ou recusada sem alterar dados.",
        "evidence": "src/app/admin; src/app/api/admin; src/services/admin.service.ts",
        "diagram": "AN09-administrar-plataforma",
    },
]


def build_activities() -> None:
    document = new_from_template("activities", "Atividades do Negócio", "Atividades do Negócio")
    document.add_heading("2. Objetivo", level=1)
    document.add_paragraph(
        "Este documento identifica as principais atividades de negócio do Stacklyst, seus participantes, decisões, exceções e resultados. "
        "Cada fluxo textual corresponde ao diagrama Mermaid apresentado na mesma subseção."
    )
    document.add_heading("3. Atividades do Negócio", level=1)
    document.add_paragraph(
        "As atividades foram organizadas do início ao fim da interação. Os números existentes no código são registrados como baseline técnico e recebem anotação quando ainda dependem de aprovação da equipe ou da professora."
    )
    for index, activity in enumerate(ACTIVITIES, 1):
        document.add_heading(f"3.{index} {activity['id']} — {activity['name']}", level=2)
        add_labeled(document, "Objetivo", activity["objective"])
        add_labeled(document, "Participantes", activity["participants"])
        add_labeled(document, "Pré-condições", activity["preconditions"])
        document.add_heading(f"3.{index}.1 Fluxo principal", level=3)
        add_numbered(document, activity["flow"])
        add_labeled(document, "Decisões", activity["decisions"])
        add_labeled(document, "Exceções", activity["exceptions"])
        add_labeled(document, "Resultado final", activity["result"])
        add_labeled(document, "Evidências de implementação", activity["evidence"])
        if activity["id"] in {"AN02", "AN07", "AN08", "AN09"}:
            add_note(document, "[ANOTAÇÃO — VALIDAR COM A EQUIPE]", "Confirmar os critérios definitivos, permissões e efeitos acadêmicos antes de tratar o fluxo como regra estável.", "yellow")
        document.add_heading(f"3.{index}.2 Diagrama de Atividades", level=3)
        image = DIAGRAMS / "imagens" / f"{activity['diagram']}.png"
        add_figure(document, image, f"Figura {index} — {activity['id']} {activity['name']}.")
        document.add_page_break()
        document.add_heading(f"3.{index}.3 Código Mermaid", level=3)
        source = (DIAGRAMS / "fontes" / "atividades" / f"{activity['diagram']}.mmd").read_text(encoding="utf-8")
        add_code_block(document, source)
        if index != len(ACTIVITIES):
            document.add_page_break()

    document.add_heading("4. Relação entre atividades e implementação", level=1)
    add_table(document, ["Atividade", "Situação", "Principais evidências"], [
        [a["id"] + " — " + a["name"],
         "Parcial" if a["id"] in {"AN02", "AN06", "AN08", "AN09"} else "Implementada no escopo auditado",
         a["evidence"]]
        for a in ACTIVITIES
    ], [2.1, 1.25, 3.4], 7.7)
    add_note(document, "[ANOTAÇÃO — VALIDAR COM A PROFESSORA]", "Confirmar se os diagramas Mermaid podem ser aceitos como notação de apoio ao UML de atividades no formato exigido pela disciplina.", "yellow")
    document.save(FINAL_FILES["activities"])


def make_rf(
    identifier: str,
    name: str,
    priority: str,
    description: str,
    actors: str,
    rules: str,
    state: str,
    evidence: str,
    inputs: str,
    processing: str,
    outputs: str,
    constraints: str,
) -> dict[str, str]:
    return {
        "id": identifier,
        "name": name,
        "priority": priority,
        "description": description,
        "actors": actors,
        "rules": rules,
        "details": f"Estado: {state}. Evidência: {evidence}.",
        "inputs": inputs,
        "processing": processing,
        "outputs": outputs,
        "constraints": constraints,
    }


RF_GROUPS = [
    ("Grupo 1 — Autenticação e Usuários", [
        make_rf("RF001", "Cadastrar usuário", "Alta", "Criar uma conta de usuário com identidade única e perfil inicial.", "Visitante", "RN017, RN025", "Implementado", "src/app/register; src/app/api/auth/register", "Nome de usuário, e-mail, senha e dados opcionais permitidos.", "Validar dados, criar identidade no provedor e registro local sem duplicidade.", "Conta criada e sessão/encaminhamento de acesso.", "E-mail e username únicos; senha tratada pelo provedor; limites de cadastro."),
        make_rf("RF002", "Autenticar usuário", "Alta", "Permitir login e criação de sessão segura.", "Usuário, Avaliador, Administrador, Empresa / Recrutador", "RN017, RN025", "Implementado", "src/app/login; src/lib/auth.ts; src/lib/supabase", "Credencial e, quando escolhido, provedor OAuth.", "Verificar identidade, emitir sessão/cookies e carregar papel autorizado.", "Sessão autenticada ou mensagem de erro.", "Credenciais não podem ser expostas; autorização deve ser revalidada no servidor."),
        make_rf("RF003", "Recuperar acesso", "Alta", "Permitir recuperação segura de conta quando a senha for esquecida.", "Usuário", "RN025", "Planejado", "nenhuma rota/tela completa de recuperação foi localizada", "E-mail ou identificador da conta.", "Solicitar token temporário ao provedor e confirmar nova credencial.", "Confirmação sem revelar a existência indevida de contas.", "[ANOTAÇÃO — DECISÃO TÉCNICA PENDENTE]: definir fluxo, expiração e telas."),
        make_rf("RF004", "Gerenciar perfil", "Alta", "Consultar e atualizar dados permitidos do perfil.", "Usuário", "RN017, RN027", "Implementado", "src/app/profile/[username]; src/app/api/profile/update", "Bio, avatar, banner, pronomes e demais campos permitidos.", "Validar propriedade, sanitizar e persistir alterações.", "Perfil atualizado e visualização pública/privada coerente.", "Somente o titular ou administrador autorizado altera os dados."),
    ]),
    ("Grupo 2 — Aprendizado", [
        make_rf("RF005", "Visualizar trilhas", "Alta", "Listar trilhas disponíveis e o progresso do usuário.", "Usuário", "RN001, RN003", "Implementado", "src/app/trails; src/lib/trailsData.ts", "Idioma/tecnologia e preferências de curso.", "Carregar catálogo, níveis, bloqueios e progresso persistido.", "Mapa de trilha com estado das etapas.", "Não expor conteúdo não publicado; funcionar sem IA."),
        make_rf("RF006", "Iniciar trilha", "Alta", "Iniciar ou retomar uma trilha selecionada.", "Usuário", "RN001, RN003", "Implementado", "src/app/trails; src/app/api/trails/course-preferences", "Trilha e nível/etapa selecionados.", "Registrar preferência e direcionar à lição elegível.", "Trilha ativa e próxima ação.", "Pré-requisitos devem ser respeitados e validados."),
        make_rf("RF007", "Acessar lição", "Alta", "Exibir uma lição em página própria com etapas interativas.", "Usuário", "RN001", "Implementado", "src/app/lesson/[lessonId]; src/lib/lessons/registry.ts", "Identificador da lição.", "Resolver conteúdo cadastrado/dinâmico e verificar disponibilidade.", "Conteúdo, progresso da sessão e exercícios.", "Lição inexistente deve gerar resposta controlada; IA é opcional."),
        make_rf("RF008", "Resolver exercício", "Alta", "Receber e avaliar respostas em diferentes tipos de interação.", "Usuário; Executor de código", "RN001, RN002", "Implementado", "src/lib/lessons/evaluators.ts; src/lib/code-runner.ts", "Resposta, código, ordem, pares, lacunas ou comando.", "Validar entrada, executar quando necessário e comparar resultado esperado.", "Acerto/erro, feedback, saída e XP elegível.", "Timeout e falha externa não podem ser tratados como acerto."),
        make_rf("RF009", "Registrar progresso", "Alta", "Persistir tentativas, checkpoints e avanço por linguagem.", "Stacklyst", "RN001, RN002, RN028", "Implementado", "QuizAttempt; LanguageTrail; src/app/api/trails/checkpoint", "Atividade concluída e identidade do usuário.", "Verificar duplicidade e atualizar tentativa, XP e trilha em transação quando aplicável.", "Progresso consolidado.", "A mesma atividade não deve conceder XP repetido."),
        make_rf("RF010", "Consultar histórico", "Média", "Apresentar histórico de evolução e atividade recente.", "Usuário", "RN027", "Implementado parcialmente", "src/app/api/quiz/weekly-activity; perfil e trilhas", "Período, usuário e trilha.", "Consolidar tentativas, XP e sequência dentro das permissões.", "Linha do tempo/resumo de progresso.", "[ANOTAÇÃO — VALIDAR COM A EQUIPE]: definir período e nível de detalhe."),
    ]),
    ("Grupo 3 — Gamificação", [
        make_rf("RF011", "Atribuir XP", "Alta", "Conceder XP apenas a atividades válidas e elegíveis.", "Stacklyst", "RN001, RN002, RN028", "Implementado", "src/services/xp.service.ts", "Usuário, linguagem, atividade e quantidade configurada.", "Validar elegibilidade e incrementar total/trilha atomicamente.", "Novo XP total e por linguagem.", "Idempotência por atividade; quantidade não deve vir livremente do cliente."),
        make_rf("RF012", "Atualizar nível", "Alta", "Recalcular o nível após mudança de XP.", "Stacklyst", "RN003", "Implementado", "src/services/xp.service.ts; src/lib/config.ts", "XP atualizado.", "Aplicar faixas de nível configuradas e detectar subida.", "Nível, faixa anterior/próxima e notificação.", "Faixas atuais precisam de validação acadêmica antes de congelamento."),
        make_rf("RF013", "Gerenciar conquistas", "Média", "Conceder e listar conquistas por critérios verificáveis.", "Usuário; Stacklyst", "RN004", "Implementado", "Badge, UserBadge; src/services/xp.service.ts", "Eventos de progresso e critérios.", "Verificar conquistas já obtidas e registrar novas sem duplicidade.", "Lista de badges e data de obtenção.", "Critérios devem ser versionados e não retroagir sem decisão."),
        make_rf("RF014", "Consultar ranking", "Alta", "Exibir classificação global ou por linguagem.", "Usuário", "RN005, RN006, RN028", "Implementado", "src/app/leaderboard; src/app/api/leaderboard", "Filtro de linguagem opcional.", "Ordenar XP e calcular posição exibida.", "Ranking com posição, usuário, XP e nível.", "Empates e paginação além do top 10 precisam ser definidos."),
        make_rf("RF015", "Gerenciar divisão competitiva", "Média", "Classificar usuários em Bronze, Prata, Ouro, Platina ou Diamante.", "Stacklyst", "RN005, RN006", "Implementado parcialmente", "getUserRankTier em src/services/duel.service.ts", "XP total do usuário.", "Aplicar limites de faixa e informar divisão.", "Divisão e nível competitivo.", "[ANOTAÇÃO — VALIDAR COM A EQUIPE]: confirmar limites e relação com ranking geral."),
    ]),
    ("Grupo 4 — Duelos", [
        make_rf("RF016", "Solicitar duelo", "Alta", "Criar convite direto ou iniciar busca automática.", "Usuário", "RN007, RN008, RN009", "Implementado", "src/app/api/duels/request", "Oponente opcional, linguagem e modo automático.", "Verificar cooldown, localizar alvo e criar convite expirável.", "Convite criado e notificação ao oponente.", "Usuário não pode desafiar sem permissão ou durante cooldown."),
        make_rf("RF017", "Realizar matchmaking", "Alta", "Selecionar oponente por compatibilidade com expansão progressiva.", "Usuário; Stacklyst", "RN007", "Implementado parcialmente", "DuelService.findMatchmakingOpponent", "Usuário, XP e linguagem.", "Buscar mesma divisão, ampliar faixa e aplicar fallback atual.", "Oponente ou mensagem de indisponibilidade.", "Não existe fila temporal persistente; algoritmo atual exige validação."),
        make_rf("RF018", "Aceitar ou rejeitar oponente", "Alta", "Permitir decisão sobre convite dentro do prazo.", "Oponente", "RN008, RN009", "Implementado", "src/app/api/duels/respond; DuelService.respondDuelRequest", "Convite e ação ACCEPT/REJECT.", "Verificar destinatário, status, expiração e atualizar contadores.", "Duelo criado ou rejeição/cooldown informados.", "Apenas o destinatário responde; convite já processado não pode ser reutilizado."),
        make_rf("RF019", "Executar desafio", "Alta", "Disponibilizar problema temporizado e receber soluções dos dois participantes.", "Desafiante; Oponente", "RN010, RN011", "Implementado parcialmente", "src/app/duels/[id]; src/app/api/duels/[id]/solution", "Código e identificador do duelo.", "Validar participante, estado e prazo; armazenar uma solução por usuário.", "Soluções registradas e estado para avaliação.", "Consequência de expiração no servidor ainda precisa ser definida."),
        make_rf("RF020", "Avaliar resultado", "Alta", "Avaliar soluções por testes automáticos, voto e/ou avaliador humano.", "Stacklyst; Avaliador", "RN012, RN013, RN016", "Implementado parcialmente", "src/app/api/evaluations; DuelEvaluation; DuelVote", "Soluções, testes, pontuações e feedback.", "Selecionar modalidade de avaliação e registrar decisão autorizada.", "Pontuações, feedback, vencedor e justificativas.", "Avaliação humana permanece obrigatória quando exigida pela modalidade."),
        make_rf("RF021", "Atualizar classificação após duelo", "Média", "Aplicar efeitos competitivos somente após resultado válido.", "Stacklyst", "RN006, RN028", "Planejado parcialmente", "duelo registra vencedor, mas não há atualização ELO consolidada", "Duelo fechado e resultado aprovado.", "Validar fraude/duplicidade e calcular efeito configurado.", "Ranking/divisão e notificações atualizados.", "[ANOTAÇÃO — REGRA PRECISA SER DEFINIDA]: fórmula e tratamento de empate."),
    ]),
    ("Grupo 5 — Avaliação Humana", [
        make_rf("RF026", "Tornar usuário elegível a avaliador", "Média", "Verificar critérios e permitir candidatura.", "Usuário; Administrador", "RN014, RN015", "Implementado", "src/services/evaluator.service.ts", "Usuário, motivação e tecnologias.", "Verificar trilha concluída/XP, duplicidade e registrar candidatura.", "Elegibilidade e candidatura pendente.", "Critérios atuais são baseline a validar; aprovação administrativa é obrigatória."),
        make_rf("RF027", "Receber avaliações pendentes", "Média", "Listar desafios com solução que necessitam revisão.", "Avaliador", "RN012, RN016, RN017", "Implementado", "GET src/app/api/evaluations", "Sessão do avaliador.", "Autorizar papel e carregar duelos pendentes.", "Fila de avaliações com soluções e contexto.", "Somente EVALUATOR ou ADMIN; dados devem ser minimizados."),
        make_rf("RF028", "Avaliar solução", "Média", "Permitir análise comparativa e parecer humano.", "Avaliador", "RN012, RN013, RN016", "Implementado", "src/app/evaluations; POST src/app/api/evaluations", "Pontuações, vencedor opcional, feedback, pontos fortes e melhorias.", "Validar limites, autorização e duelo; registrar avaliação humana.", "Parecer e resultado do duelo.", "A decisão final é registrada pelo avaliador autorizado."),
        make_rf("RF029", "Registrar decisão de avaliação", "Alta", "Persistir avaliação, fechar o duelo e notificar participantes.", "Stacklyst", "RN016, RN022, RN028", "Implementado", "EvaluatorService.submitDuelEvaluation", "Avaliação humana validada.", "Criar DuelEvaluation, definir vencedor, fechar duelo e atualizar contador.", "Decisão auditável e notificações.", "Evitar avaliações conflitantes; política de reavaliação precisa ser definida."),
    ]),
    ("Grupo 7 — Comunidade", [
        make_rf("RF030", "Criar publicação", "Alta", "Publicar conteúdo técnico no feed.", "Usuário", "RN018, RN019", "Implementado", "src/app/feed; src/app/api/posts", "Título, corpo, linguagem, código e imagem opcional.", "Validar, sanitizar, aplicar limite e persistir.", "Publicação exibida no feed.", "Autenticação, limite de conteúdo, rate limit e moderação."),
        make_rf("RF031", "Visualizar feed", "Alta", "Exibir publicações paginadas e ordenadas.", "Usuário", "RN018, RN027", "Implementado", "src/app/feed; PostService; paginação por cursor", "Filtro, limite e cursor.", "Buscar página e relacionamentos permitidos.", "Itens e próximo cursor.", "Paginação deve evitar duplicidade e carga excessiva."),
        make_rf("RF032", "Interagir com publicação", "Média", "Reagir, votar, responder, salvar e seguir conforme a ação.", "Usuário", "RN001, RN018, RN022", "Implementado", "src/app/api/posts/[id]; reactions, answers, bookmarks", "Publicação, tipo de interação e conteúdo quando aplicável.", "Validar alvo, evitar duplicidade e registrar interação.", "Contadores, resposta e notificação.", "Ação própria indevida e abuso devem ser bloqueados."),
        make_rf("RF033", "Denunciar conteúdo", "Alta", "Registrar denúncia para análise administrativa.", "Usuário; Administrador", "RN019", "Implementado", "src/app/api/posts/[id]/report; src/app/api/admin/reports", "Publicação e motivo.", "Validar usuário/alvo, impedir duplicidade e criar denúncia.", "Confirmação e item na fila administrativa.", "Denúncia não remove conteúdo automaticamente sem política definida."),
    ]),
    ("Grupo 8 — Vagas e Empresas", [
        make_rf("RF034", "Cadastrar empresa", "Média", "Criar empresa vinculada ao responsável autorizado.", "Empresa / Recrutador; Administrador", "RN017, RN020, RN027", "Implementado", "src/app/api/companies; Company", "Nome, slug, descrição, localização, site e imagem opcionais.", "Validar papel, unicidade e vínculo do proprietário.", "Empresa cadastrada.", "Somente RECRUITER/ADMIN; verificação oficial não pode ser presumida."),
        make_rf("RF035", "Publicar vaga", "Média", "Criar oportunidade com requisitos e etapas.", "Empresa / Recrutador; Administrador", "RN020, RN026", "Implementado", "POST src/app/api/jobs; JobService.createJob", "Empresa, título, descrição, nível, tecnologias, modalidade, contrato e etapas.", "Validar papel e dados, criar vaga e etapas.", "Vaga aberta e consultável.", "Responsabilidade sobre a empresa e campos obrigatórios devem ser verificados."),
        make_rf("RF036", "Consultar vagas", "Média", "Listar e filtrar oportunidades abertas.", "Usuário; Empresa / Recrutador", "RN020, RN026", "Implementado", "src/app/jobs; GET src/app/api/jobs", "Busca, nível, modalidade, contrato e tecnologia.", "Aplicar filtros e carregar empresa/etapas permitidas.", "Lista e detalhes das vagas.", "Vagas encerradas não devem aceitar candidatura."),
        make_rf("RF037", "Visualizar perfil público de desenvolvedor", "Média", "Exibir perfil público conforme configuração e autorização.", "Empresa / Recrutador; Usuário", "RN021, RN027", "Implementado parcialmente", "src/app/profile/[username]", "Username do desenvolvedor.", "Carregar apenas campos públicos e evidências permitidas.", "Perfil e indicadores públicos.", "Não expor e-mail, dados privados ou decisões automatizadas."),
        make_rf("RF038", "Consultar indicadores permitidos do usuário", "Baixa", "Apresentar trilhas, desafios e atividade pública como indicadores auxiliares.", "Empresa / Recrutador", "RN013, RN021, RN027", "Implementado parcialmente", "perfil público e modelos de progresso; não há painel analítico completo", "Usuário e permissão/consentimento aplicável.", "Selecionar indicadores públicos e registrar acesso quando necessário.", "Resumo de evidências sem decisão automática.", "Critérios de visibilidade e consentimento precisam ser aprovados."),
        make_rf("RF047", "Candidatar-se a vaga", "Média", "Registrar interesse de usuário em vaga aberta.", "Usuário", "RN026", "Implementado", "POST src/app/api/jobs/[id]/apply", "Identificador da vaga e sessão.", "Verificar vaga aberta e candidatura prévia; criar aplicação.", "Candidatura com status inicial.", "Uma candidatura por usuário/vaga."),
        make_rf("RF048", "Gerenciar candidaturas", "Média", "Atualizar etapa, status, feedback e nota técnica de candidatura.", "Empresa / Recrutador; Administrador", "RN020, RN021, RN026", "Implementado", "PATCH src/app/api/jobs/[id]; JobService.updateApplicationStage", "Candidatura, etapa, status, feedback e nota opcional.", "Autorizar responsável e atualizar processo; notificar candidato.", "Candidatura atualizada.", "Nota técnica não pode ser decisão automática baseada apenas em IA."),
    ]),
    ("Grupo 9 — Administração", [
        make_rf("RF039", "Gerenciar usuários", "Alta", "Listar usuários e alterar papéis autorizados.", "Administrador", "RN017, RN025", "Implementado", "src/app/api/admin/users; AdminService", "Busca, página, papel e alteração solicitada.", "Autorizar administrador, validar papel e persistir.", "Lista ou usuário atualizado.", "Mudança de papel deve ser auditada; impedir escalada indevida."),
        make_rf("RF040", "Gerenciar conteúdo", "Média", "Revisar conteúdo e parâmetros editoriais administráveis.", "Administrador", "RN017, RN019", "Implementado parcialmente", "admin possui denúncias e geração de quiz; CRUD editorial completo não foi localizado", "Conteúdo, decisão e justificativa.", "Validar permissão e aplicar ação suportada.", "Conteúdo atualizado/ocultado ou ação recusada.", "[ANOTAÇÃO — DECISÃO TÉCNICA PENDENTE]: definir catálogo de ações."),
        make_rf("RF041", "Gerenciar eventos", "Média", "Criar e administrar eventos comunitários.", "Administrador; Empresa / Recrutador", "RN017, RN023", "Implementado parcialmente", "src/app/api/events; EventService", "Título, descrição, tipo, datas, nível e limite.", "Validar papel/datas e persistir evento.", "Evento publicado e participável.", "Edição/cancelamento administrativo completo precisa ser validado."),
        make_rf("RF042", "Gerenciar permissões", "Alta", "Controlar papéis USER, EVALUATOR, ADMIN e RECRUITER.", "Administrador", "RN015, RN017", "Implementado", "UserRole; requireRole; admin users/evaluators", "Usuário, papel e decisão.", "Validar administrador, aplicar papel e registrar decisão.", "Permissão atualizada.", "Princípio do menor privilégio e auditoria obrigatória."),
        make_rf("RF043", "Tratar denúncias", "Alta", "Listar e decidir denúncias da comunidade.", "Administrador", "RN019", "Implementado", "src/app/api/admin/reports", "Denúncia e decisão.", "Carregar contexto, aplicar ação permitida e registrar resultado.", "Denúncia tratada e eventual alteração de conteúdo.", "Política de sanções e recurso ainda precisa ser definida."),
    ]),
    ("Grupo 10 — Notificações", [
        make_rf("RF044", "Gerar notificação", "Média", "Criar notificação para eventos relevantes.", "Stacklyst", "RN022", "Implementado", "NotificationService; NotificationType", "Destinatário, tipo, ator e recurso.", "Validar dados e persistir de forma resiliente.", "Notificação registrada.", "Falha de notificação não deve corromper a operação principal."),
        make_rf("RF045", "Visualizar notificações", "Média", "Listar, contar e marcar notificações web como lidas.", "Usuário, Avaliador, Administrador, Empresa / Recrutador", "RN022, RN027", "Implementado", "src/app/notifications; src/app/api/notifications", "Sessão, cursor/limite e ação de leitura.", "Buscar notificações do próprio usuário e atualizar leitura.", "Lista, próximo cursor e contagem não lida.", "Usuário não pode consultar notificações alheias."),
        make_rf("RF046", "Receber notificações no mobile", "Baixa", "Entregar push e visualizar informações selecionadas em aplicativo mobile.", "Usuário; Aplicação mobile", "RN022, RN027", "Planejado", "nenhum aplicativo ou serviço push localizado", "Token do dispositivo e preferências.", "Registrar dispositivo, respeitar consentimento e enviar push.", "Notificação mobile e controle de preferência.", "[ANOTAÇÃO — DECISÃO TÉCNICA PENDENTE]: plataforma, provedor, opt-in e retenção."),
    ]),
    ("Grupo 11 — Eventos", [
        make_rf("RF049", "Participar de evento", "Média", "Inscrever usuário elegível em evento comunitário.", "Usuário", "RN023", "Implementado", "POST src/app/api/events/[id]/participate; EventService.participate", "Evento e sessão do usuário.", "Validar período, nível, limite e participação prévia.", "Inscrição confirmada.", "Uma inscrição por evento; regras de desistência ainda não definidas."),
    ]),
]


RNFS = [
    ("RNF001", "Autenticação protegida", "Alta", "Segurança", "Rotas privadas devem exigir sessão verificada no servidor.", "Teste automatizado deve comprovar 401 sem sessão válida e acesso com sessão válida."),
    ("RNF002", "Autorização por papel", "Alta", "Segurança", "Operações administrativas, de avaliador e recrutador devem aplicar RBAC.", "Testes devem comprovar 403 para papel não autorizado e sucesso para papel permitido."),
    ("RNF003", "Proteção de entrada e abuso", "Alta", "Segurança", "Entradas devem ser validadas e endpoints críticos devem limitar frequência.", "Schemas rejeitam payload inválido e testes verificam limite configurado sem expor segredo."),
    ("RNF004", "Segredos server-only", "Alta", "Segurança", "Credenciais de banco, IA e serviço não podem ser enviadas ao cliente.", "Build e auditoria estática não contêm segredo em variáveis NEXT_PUBLIC nem no bundle cliente."),
    ("RNF005", "Separação público/privado", "Alta", "Privacidade", "Perfil e recrutamento devem retornar apenas campos autorizados.", "Teste com usuário/empresa confirma ausência de e-mail, token e dados privados nas respostas públicas."),
    ("RNF006", "Consentimento de recrutamento", "Alta", "Privacidade", "Indicadores usados por empresas dependem de regra de visibilidade e consentimento.", "[ANOTAÇÃO — DEFINIR MÉTRICA]: validar percentual/cobertura e telas de opt-in/opt-out."),
    ("RNF007", "Tempo de resposta das APIs", "Média", "Desempenho", "APIs interativas devem atender meta no percentil definido sob carga representativa.", "P95 ≤ [ANOTAÇÃO — DEFINIR MÉTRICA] para [ANOTAÇÃO — DEFINIR CARGA]."),
    ("RNF008", "Carregamento de página", "Média", "Desempenho", "Páginas principais devem ter meta de carregamento em dispositivo/rede definidos.", "LCP ≤ [ANOTAÇÃO — DEFINIR MÉTRICA] no cenário [ANOTAÇÃO — DEFINIR CENÁRIO]."),
    ("RNF009", "Execução limitada", "Alta", "Desempenho", "Execução de código e IA devem possuir timeout e limite de payload.", "Teste provoca timeout e confirma resposta controlada sem bloquear o servidor além do limite configurado."),
    ("RNF010", "Responsividade", "Alta", "Usabilidade", "Fluxos principais devem funcionar em desktop e largura mobile web.", "Teste visual em larguras [ANOTAÇÃO — DEFINIR VIEWPORTS] sem perda de ação essencial."),
    ("RNF011", "Acessibilidade básica", "Média", "Usabilidade", "Controles devem possuir nome acessível, foco e suporte a teclado.", "Auditoria automática e roteiro manual atingem [ANOTAÇÃO — DEFINIR MÉTRICA] sem bloqueador crítico."),
    ("RNF012", "Degradação de serviço externo", "Alta", "Disponibilidade", "Falha de IA, cache ou notificação não deve interromper fluxos essenciais.", "Teste simula indisponibilidade e confirma fallback/erro controlado sem perda de dados."),
    ("RNF013", "Meta de disponibilidade", "Média", "Disponibilidade", "A plataforma deve possuir meta mensal e rotina de monitoramento.", "Disponibilidade ≥ [ANOTAÇÃO — DEFINIR MÉTRICA] e alertas para falhas críticas."),
    ("RNF014", "Consistência de XP", "Alta", "Confiabilidade", "XP e progresso devem ser atualizados atomicamente e sem duplicidade.", "Teste concorrente comprova uma recompensa por atividade e rollback integral em falha."),
    ("RNF015", "Consistência de duelo e avaliação", "Alta", "Confiabilidade", "Estados de convite, duelo e avaliação devem impedir transições inválidas.", "Testes cobrem convite expirado, dupla resposta, dupla solução e fechamento conflitante."),
    ("RNF016", "Crescimento de listas", "Média", "Escalabilidade", "Feed, notificações e históricos devem usar paginação/índices.", "Teste com [ANOTAÇÃO — DEFINIR VOLUME] mantém P95 dentro da meta RNF007."),
    ("RNF017", "Navegadores suportados", "Média", "Compatibilidade", "A aplicação web deve funcionar nos navegadores definidos pela equipe.", "[ANOTAÇÃO — DEFINIR MÉTRICA]: lista e versões mínimas de navegadores."),
    ("RNF018", "Compatibilidade mobile", "Baixa", "Compatibilidade", "O futuro aplicativo deve compartilhar contratos de API sem quebrar a web.", "Testes de contrato executam nos clientes definidos após decisão técnica."),
    ("RNF019", "Modularização", "Média", "Manutenibilidade", "Domínio, APIs, UI e integrações externas devem permanecer separados.", "Lint, typecheck e testes passam; revisão não identifica dependência circular crítica."),
    ("RNF020", "Cobertura de qualidade", "Média", "Manutenibilidade", "Mudanças em regras críticas devem incluir teste e documentação.", "CI executa typecheck, lint e testes; cobertura mínima: [ANOTAÇÃO — DEFINIR MÉTRICA]."),
    ("RNF021", "IA auxiliar e segura", "Alta", "IA", "IA deve exibir natureza auxiliar, validar saída e nunca substituir decisão humana obrigatória.", "Teste confirma que avaliação/moderação humana não é finalizada somente por resposta de IA."),
    ("RNF022", "Limites e fallback da IA", "Alta", "IA", "Chamadas devem limitar corpo, contexto, tempo e frequência e apresentar fallback.", "Testes cobrem payload excedido, quota, timeout, provedor ausente e resposta inválida."),
    ("RNF023", "Minimização de contexto da IA", "Alta", "IA", "Somente contexto necessário e autorizado deve ser enviado.", "Auditoria comprova ausência de segredo e respeita limites configurados de mensagens/caracteres."),
    ("RNF024", "Auditoria de ações sensíveis", "Alta", "Auditoria", "Alterações de papel, avaliações e ações administrativas devem ser rastreáveis.", "Registro contém ator, ação, alvo, instante e resultado para 100% das operações definidas."),
    ("RNF025", "Retenção de auditoria", "Média", "Auditoria", "Logs sensíveis devem possuir retenção e acesso controlados.", "[ANOTAÇÃO — DEFINIR MÉTRICA]: prazo, responsável, anonimização e descarte."),
]


BUSINESS_RULES = [
    ("RN001", "Elegibilidade de XP", "XP somente é concedido após atividade válida e confirmada; repetir a mesma tentativa não gera nova recompensa.", "Implementado", "QuizAttempt e XpService"),
    ("RN002", "Valores de XP", "Valores são configurados pelo servidor. Baseline atual: quiz correto 15 XP e checkpoint 50 XP.", "Validar", "src/lib/config.ts; checkpoint route"),
    ("RN003", "Níveis", "Níveis são calculados por faixas de XP. Baseline atual: 500, 800, 1.100, 1.500 e 2.000 XP, com progressão incremental após o nível 6.", "Validar", "src/services/xp.service.ts; src/lib/config.ts"),
    ("RN004", "Sequência e conquistas", "Atividade em dias consecutivos atualiza sequência; conquistas não podem ser duplicadas. Critérios atuais incluem 7 e 30 dias.", "Validar", "src/lib/streak.ts; XpService"),
    ("RN005", "Divisões competitivas", "As divisões são Bronze, Prata, Ouro, Platina e Diamante. Limites atuais: 0, 500, 1.200, 2.500 e 5.000 XP.", "Validar", "getUserRankTier"),
    ("RN006", "Ordenação do ranking", "Ranking global usa XP total; ranking por linguagem usa XP da trilha. Critério de desempate precisa ser definido.", "Parcial", "src/app/api/leaderboard"),
    ("RN007", "Matchmaking progressivo", "Buscar mesma divisão; sem candidato, ampliar para ±1.000 XP; por fim, usar usuário disponível. Essa regra diverge da proposta de buscar somente faixa superior.", "Decisão necessária", "DuelService.findMatchmakingOpponent"),
    ("RN008", "Prazo do convite", "Baseline atual: convite expira em 30 segundos.", "Validar", "DUEL_REQUEST_TIMEOUT_SECONDS"),
    ("RN009", "Rejeição e cooldown", "Baseline atual: 3 rejeições consecutivas aplicam cooldown de 5 minutos e zeram o contador.", "Validar", "MAX_DUEL_REJECTIONS; DUEL_COOLDOWN_MINUTES"),
    ("RN010", "Tempo do desafio", "Baseline atual do duelo aceito: 900 segundos (15 minutos).", "Validar", "time_limit_seconds em DuelService"),
    ("RN011", "Expiração da execução", "Após o prazo, o sistema deve impedir solução e registrar não conclusão ou retorno ao estágio correspondente.", "Regra precisa ser definida", "não há política server-side completa localizada"),
    ("RN012", "Modalidade de avaliação", "O tipo de desafio determina avaliação automática, por voto e/ou humana.", "Validar", "EvaluationType; DuelVote; evaluations API"),
    ("RN013", "IA não autoritativa", "IA fornece sugestão; decisões humanas obrigatórias e efeitos competitivos não podem depender somente dela.", "Oficial no escopo", "prompts e requisito do projeto"),
    ("RN014", "Elegibilidade de avaliador", "Baseline atual: concluir ao menos uma trilha completa ou possuir 1.000 XP.", "Validar", "EvaluatorService.checkEligibility"),
    ("RN015", "Aprovação de avaliador", "A elegibilidade permite candidatura, mas somente o administrador atribui o papel EVALUATOR.", "Implementado", "EvaluatorService.reviewApplication"),
    ("RN016", "Parecer humano", "Pontuações atuais variam de 0 a 100 e exigem feedback com pelo menos 5 caracteres; conflito/empate precisa de política.", "Validar", "evaluationSchema"),
    ("RN017", "Controle por papel", "USER usa funções gerais; EVALUATOR avalia; ADMIN administra; RECRUITER gerencia empresa, vagas e eventos autorizados.", "Implementado", "UserRole; requireRole"),
    ("RN018", "Publicação comunitária", "Publicação exige autenticação, validação de conteúdo e limites. Baseline atual: corpo de até 5.000 caracteres e limite de criação configurado.", "Validar", "src/lib/config.ts; posts API"),
    ("RN019", "Moderação", "Denúncia deve ser registrada uma vez por usuário/publicação e analisada por administrador; sanções e recurso ainda dependem de política.", "Parcial", "Report; admin reports"),
    ("RN020", "Publicação de vaga", "Somente RECRUITER ou ADMIN publica vaga em empresa autorizada; responsabilidade pela empresa deve ser verificada.", "Implementado parcialmente", "jobs API; companies API"),
    ("RN021", "Acesso de empresas", "Empresas acessam somente dados públicos/permitidos; desempenho é indicador auxiliar e não decisão automática.", "Decisão necessária", "perfil público; requisito do projeto"),
    ("RN022", "Notificações", "Eventos relevantes geram notificação ao destinatário; falha de notificação não desfaz a operação principal.", "Implementado na web", "NotificationService"),
    ("RN023", "Participação em evento", "Usuário participa uma vez, respeitando período, nível mínimo e capacidade quando definidos.", "Implementado", "EventService.participate"),
    ("RN024", "Limites e fallback de IA", "IA deve respeitar quotas, tamanho, timeout e disponibilidade; falha aciona erro controlado ou conteúdo curado.", "Implementado parcialmente", "src/lib/ratelimit.ts; src/lib/ai"),
    ("RN025", "Identidade e sessão", "E-mail e username são únicos; rotas protegidas exigem sessão verificada.", "Implementado", "User; auth/register; requireAuth"),
    ("RN026", "Candidatura a vaga", "Usuário pode candidatar-se uma vez a vaga aberta; empresa atualiza etapa/status e o candidato é notificado.", "Implementado", "JobApplication; JobService"),
    ("RN027", "Privacidade do perfil", "A equipe deve definir campos públicos, privados e consentimento para recrutamento.", "Decisão necessária", "[ANOTAÇÃO — REGRA PRECISA SER DEFINIDA]"),
    ("RN028", "Integridade competitiva", "XP, ranking, duelo e avaliação devem rejeitar duplicidade, manipulação pelo cliente e transições inválidas.", "Parcial", "transações/uniques existentes; auditoria a ampliar"),
]


def add_requirement_table(document: Document, requirement: dict[str, str]) -> None:
    rows = [
        ["Nome", requirement["name"]],
        ["Identificador", requirement["id"]],
        ["Prioridade", requirement["priority"]],
        ["Descrição", requirement["description"]],
        ["Atores envolvidos", requirement["actors"]],
        ["Regra de negócio relacionada", requirement["rules"]],
        ["Detalhes da implementação prevista", requirement["details"]],
        ["Entradas", requirement["inputs"]],
        ["Processamento", requirement["processing"]],
        ["Saídas", requirement["outputs"]],
        ["Restrições", requirement["constraints"]],
    ]
    add_table(document, [f"{requirement['id']} — Especificação", "Conteúdo"], rows, [1.7, 5.05], 7.4)


def build_requirements() -> None:
    document = new_from_template("requirements", "Documento de Requisitos", "Documento de Requisitos do Sistema")
    document.add_heading("2. Objetivo", level=1)
    document.add_paragraph(
        "Este documento identifica, organiza e formaliza os requisitos funcionais, os requisitos não funcionais e as regras de negócio do Stacklyst. "
        "Também registra prioridade, atores, entradas, processamento, saídas, restrições, estado de implementação e evidências no repositório."
    )
    document.add_heading("3. Técnicas Utilizadas na Elucidação de Requisitos", level=1)
    document.add_paragraph("Não foram fornecidas atas, entrevistas ou questionários que comprovem técnicas já executadas. As técnicas abaixo permanecem como itens a confirmar:")
    for technique in [
        "Brainstorming da equipe", "Análise de plataformas similares", "Prototipação", "Reuniões acadêmicas",
        "Validação com professora/orientadora", "Análise das necessidades dos usuários",
    ]:
        add_note(document, "[ANOTAÇÃO — CONFIRMAR SE ESTA TÉCNICA FOI UTILIZADA]", technique, "yellow")
    document.add_paragraph("Nesta versão, a principal fonte comprovada foi a análise documental dos quatro templates, das instruções do projeto e do repositório atual do Stacklyst.")

    document.add_heading("4. Requisitos Funcionais", level=1)
    document.add_paragraph("Os requisitos foram agrupados por módulo. O estado informado não substitui testes de aceitação; ele indica apenas a evidência localizada no código auditado.")
    section_index = 1
    all_requirements: list[dict[str, str]] = []
    for group_name, requirements in RF_GROUPS:
        document.add_heading(f"4.{section_index} {group_name}", level=2)
        for item_index, requirement in enumerate(requirements, 1):
            document.add_heading(f"4.{section_index}.{item_index} {requirement['id']} — {requirement['name']}", level=3)
            add_requirement_table(document, requirement)
            all_requirements.append(requirement)
        section_index += 1

    document.add_heading("5. Requisitos Não Funcionais", level=1)
    document.add_paragraph("Critérios sem valor aprovado usam a anotação solicitada e não devem ser tratados como metas oficiais.")
    add_table(document, ["Identificador", "Nome", "Prioridade", "Categoria", "Descrição", "Critério verificável de aceitação"], [list(r) for r in RNFS], [0.62, 1.1, 0.62, 0.78, 1.65, 2.0], 6.5)

    document.add_heading("6. Regras de Negócio", level=1)
    add_table(document, ["ID", "Regra", "Descrição", "Situação", "Evidência"], [list(r) for r in BUSINESS_RULES], [0.45, 1.15, 3.05, 0.85, 1.25], 6.7)
    add_note(document, "[ANOTAÇÃO — REGRA PRECISA SER DEFINIDA]", "Itens com situação Validar, Decisão necessária, Parcial ou Regra precisa ser definida exigem aprovação antes do baseline oficial.", "yellow")

    document.add_heading("7. Matriz de relacionamento Requisitos Funcionais × Regras de Negócio", level=1)
    add_table(document, ["Requisito Funcional", "Regra(s) de Negócio Relacionada(s)"], [[f"{r['id']} — {r['name']}", r["rules"]] for r in all_requirements], [4.55, 2.2], 8)
    document.save(FINAL_FILES["requirements"])


def make_uc(
    identifier: str,
    name: str,
    use_type: str,
    initiators: str,
    secondary: str,
    description: str,
    preconditions: str,
    postconditions: str,
    inputs: str,
    outputs: str,
    flow: list[str],
    alternatives: list[str],
    prototype: str,
) -> dict[str, object]:
    return {
        "id": identifier,
        "name": name,
        "type": use_type,
        "initiators": initiators,
        "secondary": secondary,
        "description": description,
        "preconditions": preconditions,
        "postconditions": postconditions,
        "inputs": inputs,
        "outputs": outputs,
        "flow": flow,
        "alternatives": alternatives,
        "prototype": prototype,
    }


USE_CASES = [
    make_uc("UC001", "Cadastrar usuário", "Condução", "Usuário / Estudante / Desenvolvedor", "Serviço de autenticação",
        "Criar uma conta individual para acessar funções autenticadas.", "Visitante não autenticado; e-mail e username disponíveis.",
        "Identidade e perfil inicial criados; usuário pode autenticar-se.", "Username, e-mail, senha e campos opcionais permitidos.",
        "Confirmação de cadastro ou erros por campo.",
        ["O visitante abre a tela de cadastro e informa os dados solicitados.", "O Stacklyst valida formato, força da senha e unicidade.", "O sistema solicita ao serviço de autenticação a criação da identidade.", "Após confirmação, o Stacklyst cria o registro local e perfil inicial.", "O sistema confirma o cadastro e direciona o usuário ao acesso apropriado."],
        ["A1 — Dado inválido: o sistema destaca o campo sem criar a conta.", "A2 — E-mail ou username existente: o sistema recusa duplicidade.", "A3 — Serviço indisponível: o sistema informa falha temporária sem expor detalhes."],
        "Tela de cadastro (/register)"),
    make_uc("UC002", "Realizar login", "Condução", "Usuário, Avaliador, Administrador, Empresa / Recrutador", "Serviço de autenticação",
        "Autenticar um ator e iniciar sessão com o papel correspondente.", "Conta existente e não bloqueada.",
        "Sessão válida criada; rotas autorizadas ficam disponíveis.", "Credencial ou provedor OAuth.", "Sessão, perfil básico e redirecionamento.",
        ["O ator abre a tela de login e informa a credencial.", "O Stacklyst envia a solicitação ao serviço de autenticação.", "O serviço verifica a identidade e retorna o resultado.", "O Stacklyst cria/atualiza cookies de sessão e carrega o papel.", "O sistema redireciona para a área permitida."],
        ["A1 — Credencial inválida: o sistema mantém a sessão encerrada.", "A2 — Provedor OAuth não configurado: o sistema informa indisponibilidade.", "A3 — Conexão temporária: o sistema permite nova tentativa sem revelar segredo."],
        "Tela de login (/login)"),
    make_uc("UC003", "Gerenciar perfil", "Configuração", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Consultar e atualizar dados permitidos do próprio perfil.", "Usuário autenticado.",
        "Campos válidos persistidos e perfil público atualizado.", "Bio, avatar, banner, pronomes e demais campos suportados.", "Perfil atualizado ou erros de validação.",
        ["O usuário acessa o próprio perfil ou as configurações.", "O Stacklyst exibe os dados editáveis e a visualização atual.", "O usuário altera campos e confirma.", "O sistema valida propriedade, formato e limites e persiste os dados.", "O Stacklyst recarrega o perfil e confirma a atualização."],
        ["A1 — Usuário sem permissão: o sistema nega alteração em perfil alheio.", "A2 — Campo inválido: o sistema preserva dados anteriores.", "A3 — Upload indisponível: o sistema mantém a imagem anterior."],
        "uc003-uc007-perfil-progresso-atual.png"),
    make_uc("UC004", "Iniciar trilha", "Condução", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Selecionar uma trilha e iniciar ou retomar o percurso.", "Usuário autenticado; trilha disponível.",
        "Preferência registrada e lição elegível apresentada.", "Linguagem/trilha e etapa selecionada.", "Mapa atualizado e acesso à próxima lição.",
        ["O usuário abre Trilhas e consulta as opções.", "O Stacklyst carrega progresso, níveis e bloqueios.", "O usuário seleciona a trilha desejada.", "O sistema registra a preferência e determina a etapa elegível.", "O Stacklyst abre a lição ou apresenta o próximo passo."],
        ["A1 — Trilha indisponível: o sistema impede o início e informa o motivo.", "A2 — Pré-requisito não cumprido: o sistema mantém a etapa bloqueada.", "A3 — Progresso não carregado: o sistema não sobrescreve dados existentes."],
        "uc004-trilhas-atual.png"),
    make_uc("UC005", "Acessar lição", "Condução", "Usuário / Estudante / Desenvolvedor", "Stacklyst; Serviço de IA opcional",
        "Exibir conteúdo e etapas interativas de uma lição.", "Usuário autenticado; lição existente e elegível.",
        "Sessão de lição iniciada; exercícios podem ser realizados.", "Identificador da lição.", "Conteúdo, etapas e progresso da sessão.",
        ["O usuário escolhe uma lição liberada na trilha.", "O Stacklyst resolve o identificador no registro de lições.", "O sistema verifica conteúdo e estado da etapa.", "A página apresenta explicação, exemplos e interação correspondente.", "O usuário navega pelas etapas e decide resolver o exercício ou pedir ajuda opcional."],
        ["A1 — Lição inexistente: o sistema apresenta estado de não encontrada.", "A2 — Lição bloqueada: o sistema orienta o pré-requisito.", "A3 — IA indisponível: a lição permanece utilizável sem assistência."],
        "Tela da lição (/lesson/[lessonId])"),
    make_uc("UC006", "Resolver exercício", "Condução", "Usuário / Estudante / Desenvolvedor", "Executor de código; Stacklyst",
        "Avaliar resposta em um dos tipos interativos suportados.", "Exercício carregado e usuário autenticado para persistência.",
        "Tentativa, feedback e progresso elegível registrados.", "Resposta, código, pares, ordem, lacunas, comando ou blocos.", "Resultado, feedback, saída e XP quando aplicável.",
        ["O usuário preenche ou monta a solução do exercício.", "O Stacklyst valida se a interação está completa.", "Quando há código, o sistema executa em serviço controlado e recebe saída/testes.", "O Stacklyst compara o resultado com o esperado e produz feedback.", "No acerto inédito, o sistema registra a tentativa e atualiza XP/progresso."],
        ["A1 — Dado inválido: o sistema solicita correção antes da avaliação.", "A2 — Executor indisponível ou timeout: o sistema não marca acerto nem concede XP.", "A3 — Resposta incorreta: o sistema exibe feedback e permite nova tentativa."],
        "Tela do exercício dentro da lição"),
    make_uc("UC007", "Consultar progresso", "Análise", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Apresentar XP, nível, sequência, trilhas, tentativas e conquistas permitidas.", "Usuário autenticado ou perfil público acessível.",
        "Nenhuma alteração obrigatória; visão consolidada exibida.", "Usuário, período e trilha opcionais.", "Indicadores e histórico permitido.",
        ["O usuário abre perfil, trilha ou visão de atividade.", "O Stacklyst identifica o titular e as permissões de visualização.", "O sistema consulta XP, trilhas, sequência, tentativas e badges.", "Os dados são consolidados sem expor campos privados.", "O Stacklyst apresenta evolução atual e próximos marcos."],
        ["A1 — Perfil privado ou dado não permitido: o sistema oculta o campo.", "A2 — Histórico indisponível: o sistema mostra o último estado confirmado.", "A3 — Período sem atividade: o sistema retorna estado vazio claro."],
        "uc003-uc007-perfil-progresso-atual.png"),
    make_uc("UC008", "Participar de duelo", "Condução", "Usuário / Estudante / Desenvolvedor", "Oponente; Avaliador; Executor de código",
        "Executar um desafio competitivo entre dois usuários.", "Participantes autenticados, convite aceito e sem cooldown.",
        "Duelo fechado, não concluído ou encaminhado para avaliação.", "Linguagem, convite e soluções.", "Estado, resultado, feedback e notificações.",
        ["O desafiante inicia duelo direto ou busca automática.", "Após aceite, o Stacklyst cria problema e prazo do duelo.", "Os dois participantes analisam o desafio e enviam uma solução.", "O sistema recebe as soluções e executa/encaminha a avaliação definida.", "O Stacklyst fecha o duelo e apresenta o resultado aos participantes."],
        ["A1 — Nenhum adversário: o sistema informa indisponibilidade.", "A2 — Usuário rejeita ou convite expira: o duelo não é criado e a regra de rejeição é aplicada.", "A3 — Desafio expira: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA] para resultado/retorno.", "A4 — Avaliação humana necessária: o sistema encaminha à fila sem decidir somente por IA."],
        "uc008-uc009-duelos-atual.png"),
    make_uc("UC009", "Realizar matchmaking", "Condução", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Localizar oponente compatível para duelo automático.", "Usuário autenticado e apto a duelar.",
        "Oponente selecionado e convite enviado, ou busca encerrada sem candidato.", "Usuário, XP e linguagem.", "Oponente ou indisponibilidade.",
        ["O usuário solicita busca automática.", "O Stacklyst verifica cooldown e calcula a divisão atual.", "O sistema procura candidato na mesma divisão.", "Sem candidato, amplia a faixa conforme baseline e aplica fallback configurado.", "Ao localizar candidato, o Stacklyst cria e envia o convite."],
        ["A1 — Usuário em cooldown: a busca é recusada com tempo restante.", "A2 — Nenhum candidato: o sistema orienta nova tentativa.", "A3 — Candidato torna-se indisponível: o sistema não cria duelo inválido."],
        "uc008-uc009-duelos-atual.png"),
    make_uc("UC010", "Avaliar desafio", "Análise", "Avaliador ou Administrador autorizado", "Serviço de IA auxiliar; Executor de código",
        "Determinar resultado de desafio com modalidade de avaliação adequada.", "Duelo com soluções e ator autorizado quando houver decisão humana.",
        "Avaliação registrada e duelo apto a ser fechado.", "Soluções, testes, votos, pontuações e feedback.", "Parecer, vencedor e justificativas.",
        ["O ator autorizado abre um desafio com soluções.", "O Stacklyst apresenta código, resultados de testes e contexto permitido.", "O avaliador compara correção, qualidade e critérios definidos; uma sugestão de IA pode ser consultada.", "O ator informa pontuações, vencedor e feedback.", "O sistema valida e registra a modalidade e a decisão."],
        ["A1 — Dados incompletos: o sistema recusa o parecer.", "A2 — Usuário sem papel: o sistema retorna acesso negado.", "A3 — IA indisponível: a avaliação humana continua sem prejuízo.", "A4 — Conflito/empate: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA]."],
        "Tela de avaliação de desafio (/evaluations)"),
    make_uc("UC011", "Consultar ranking", "Análise", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Exibir classificação global ou por linguagem.", "Dados de XP disponíveis.", "Ranking exibido sem alterar pontuação.", "Filtro de linguagem opcional.", "Posição, usuário, XP, nível e divisão quando aplicável.",
        ["O usuário abre o ranking.", "O Stacklyst recebe o filtro global ou por linguagem.", "O sistema consulta os maiores valores de XP no escopo.", "A posição e o nível exibido são calculados.", "O Stacklyst apresenta a classificação e a posição disponível."],
        ["A1 — Filtro inválido: o sistema usa opção válida ou informa erro.", "A2 — Ranking vazio: o sistema exibe estado sem participantes.", "A3 — Empate: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA]."],
        "uc011-ranking-atual.png"),
    make_uc("UC013", "Criar publicação", "Condução", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Publicar dúvida, conteúdo ou código no feed.", "Usuário autenticado e dentro do limite de publicação.",
        "Publicação persistida e exibida no feed.", "Título, corpo, linguagem, código e imagem opcional.", "Publicação criada ou erros de validação.",
        ["O usuário abre o compositor do feed.", "O usuário informa o conteúdo e confirma a publicação.", "O Stacklyst valida autenticação, tamanho, formato e limite.", "O sistema persiste a publicação e seus metadados.", "O feed é atualizado e a publicação fica disponível para interação."],
        ["A1 — Conteúdo inválido ou vazio: o sistema solicita correção.", "A2 — Limite excedido: o sistema recusa temporariamente.", "A3 — Upload falha: o usuário pode corrigir ou publicar sem a mídia, se permitido."],
        "uc013-uc014-feed-atual.png"),
    make_uc("UC014", "Interagir com publicação", "Condução", "Usuário / Estudante / Desenvolvedor", "Autor da publicação; Stacklyst",
        "Reagir, votar, responder, salvar ou denunciar uma publicação.", "Usuário autenticado; publicação existente.",
        "Interação registrada e contadores/notificações atualizados.", "Publicação, ação, reação ou resposta.", "Novo estado da interação e eventual notificação.",
        ["O usuário seleciona uma publicação no feed.", "O usuário escolhe a interação permitida e informa dados adicionais quando necessário.", "O Stacklyst valida alvo, sessão e duplicidade.", "O sistema registra ou alterna a interação e atualiza contadores.", "O autor recebe notificação quando a regra prevê."],
        ["A1 — Publicação removida: o sistema interrompe a ação.", "A2 — Ação duplicada: o sistema mantém unicidade ou alterna o estado.", "A3 — Conteúdo denunciado: o sistema cria item para moderação."],
        "uc013-uc014-feed-atual.png"),
    make_uc("UC015", "Gerenciar evento", "Configuração", "Administrador; Empresa / Recrutador", "Stacklyst",
        "Criar e administrar evento ou desafio comunitário autorizado.", "Ator autenticado com papel ADMIN ou RECRUITER.",
        "Evento criado/atualizado com regras e período válidos.", "Título, descrição, tipo, datas, nível, capacidade e recompensa.", "Evento e status.",
        ["O ator acessa o módulo de eventos e solicita criação.", "O Stacklyst confirma o papel autorizado.", "O ator informa regras, período e capacidade.", "O sistema valida datas, limites e vínculo com empresa quando aplicável.", "O Stacklyst registra e disponibiliza o evento conforme o status."],
        ["A1 — Sem permissão: o sistema nega a operação.", "A2 — Data ou limite inválido: o sistema não publica.", "A3 — Cancelamento/edição: [ANOTAÇÃO — DECISÃO TÉCNICA PENDENTE] para efeitos nas inscrições."],
        "Tela de administração de evento (/events ou painel administrativo)"),
    make_uc("UC016", "Avaliar solução como avaliador", "Análise", "Avaliador", "Administrador; Serviço de IA auxiliar",
        "Emitir parecer humano sobre soluções de duelo.", "Avaliador aprovado; duelo pendente com solução.",
        "DuelEvaluation humana criada, duelo fechado e participantes notificados.", "Pontuações, vencedor, feedback, pontos fortes e melhorias.", "Parecer persistido e resultado.",
        ["O avaliador abre a fila de avaliações pendentes.", "O Stacklyst autoriza o papel e apresenta as soluções.", "O avaliador revisa código, testes e critérios e pode consultar sugestão auxiliar.", "O avaliador informa pontuações e feedback explicativo.", "O sistema valida, registra o parecer, fecha o duelo e notifica."],
        ["A1 — Papel ausente: o sistema retorna acesso negado.", "A2 — Duelo já fechado: o sistema impede nova decisão conflitante.", "A3 — Feedback/pontuação inválidos: o sistema solicita correção."],
        "Tela de fila e formulário de avaliação (/evaluations)"),
    make_uc("UC017", "Publicar vaga", "Configuração", "Empresa / Recrutador; Administrador", "Stacklyst",
        "Criar uma oportunidade vinculada a empresa autorizada.", "Ator com papel RECRUITER/ADMIN e empresa cadastrada.",
        "Vaga aberta com etapas de recrutamento.", "Empresa, título, descrição, nível, tecnologias, modalidade, contrato e etapas.", "Vaga publicada ou erros por campo.",
        ["O recrutador seleciona a empresa sob sua responsabilidade.", "O ator informa dados da vaga e etapas.", "O Stacklyst valida papel, empresa e campos obrigatórios.", "O sistema cria a vaga e as etapas em ordem.", "A oportunidade é disponibilizada para consulta e candidatura."],
        ["A1 — Empresa não autorizada: o sistema nega publicação.", "A2 — Vaga inválida: o sistema destaca campos e não publica.", "A3 — Salário omitido: o sistema aplica a política definida sem inventar valor."],
        "Tela de recrutador para criação de vaga (/recruiter)"),
    make_uc("UC018", "Consultar vagas", "Análise", "Usuário / Estudante / Desenvolvedor; Empresa / Recrutador", "Stacklyst",
        "Listar e detalhar oportunidades disponíveis.", "Acesso ao módulo; vagas cadastradas.", "Lista/detalhe exibido sem alteração obrigatória.", "Busca, tecnologia, nível, modalidade e contrato.", "Vagas filtradas, empresa, requisitos e etapas.",
        ["O ator abre a listagem de vagas.", "O Stacklyst carrega vagas abertas e filtros.", "O ator informa critérios ou seleciona uma oportunidade.", "O sistema aplica filtros e carrega os detalhes permitidos.", "A vaga, empresa e opção de candidatura são apresentadas."],
        ["A1 — Nenhuma vaga: o sistema exibe estado vazio.", "A2 — Vaga encerrada: o sistema exibe status e desabilita candidatura.", "A3 — Filtro inválido: o sistema solicita correção."],
        "Tela de vagas (/jobs e /jobs/[id])"),
    make_uc("UC019", "Administrar usuários", "Configuração", "Administrador", "Stacklyst",
        "Pesquisar usuários e alterar papéis autorizados.", "Administrador autenticado.",
        "Papel atualizado e ação auditável, ou operação recusada.", "Busca, paginação, usuário e novo papel.", "Lista de usuários ou registro atualizado.",
        ["O administrador abre o painel de usuários.", "O Stacklyst valida a sessão administrativa e lista os registros.", "O administrador pesquisa e seleciona o usuário.", "O administrador escolhe o novo papel e confirma.", "O sistema valida a transição, atualiza e registra a ação."],
        ["A1 — Sem permissão: o sistema nega acesso.", "A2 — Papel inválido: o sistema não altera o usuário.", "A3 — Conflito administrativo: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA] para impedir perda do último administrador."],
        "Tela administrativa de usuários (/admin)"),
    make_uc("UC020", "Gerenciar denúncias", "Análise", "Administrador", "Usuário denunciante; Autor do conteúdo",
        "Analisar denúncia e aplicar ação de moderação permitida.", "Administrador autenticado; denúncia registrada.",
        "Denúncia tratada e ação registrada.", "Denúncia, conteúdo, contexto, decisão e justificativa.", "Resultado de moderação e eventual notificação.",
        ["O administrador abre a fila de denúncias.", "O Stacklyst apresenta denúncia, conteúdo e dados mínimos necessários.", "O administrador analisa motivo e contexto.", "O administrador seleciona decisão permitida e registra justificativa.", "O sistema aplica a ação, registra auditoria e notifica quando previsto."],
        ["A1 — Denúncia improcedente: o sistema encerra sem remover conteúdo.", "A2 — Conteúdo já removido: o sistema registra o estado encontrado.", "A3 — Sanção ou recurso: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA]."],
        "Tela administrativa de denúncias (/admin)"),
    make_uc("UC021", "Visualizar notificações", "Análise", "Usuário, Avaliador, Administrador, Empresa / Recrutador", "Stacklyst",
        "Consultar eventos dirigidos à própria conta e marcar como lidos.", "Ator autenticado.",
        "Notificações exibidas e leitura atualizada quando solicitada.", "Cursor, limite e ação de leitura.", "Lista, próximo cursor e contagem não lida.",
        ["O ator abre a central de notificações.", "O Stacklyst identifica o usuário da sessão.", "O sistema busca somente notificações do destinatário com paginação.", "O ator abre itens ou solicita marcar como lidos.", "O Stacklyst atualiza leitura e a contagem não lida."],
        ["A1 — Conexão indisponível: o sistema preserva o estado e permite nova tentativa.", "A2 — Notificação removida: o sistema ignora o item sem acessar recurso inexistente.", "A3 — Mobile: fluxo planejado, sem implementação confirmada."],
        "Tela de notificações (/notifications)"),
    make_uc("UC022", "Candidatar-se a vaga", "Condução", "Usuário / Estudante / Desenvolvedor", "Empresa / Recrutador; Stacklyst",
        "Registrar candidatura de usuário em vaga aberta.", "Usuário autenticado; vaga aberta; nenhuma candidatura anterior.",
        "Candidatura criada no status inicial.", "Identificador da vaga e sessão.", "Confirmação e situação da candidatura.",
        ["O usuário consulta o detalhe de uma vaga aberta.", "O usuário seleciona a opção de candidatura.", "O Stacklyst verifica sessão, status da vaga e duplicidade.", "O sistema cria a candidatura vinculada ao usuário e à vaga.", "O Stacklyst confirma e passa a exibir a situação no processo."],
        ["A1 — Vaga encerrada: o sistema recusa candidatura.", "A2 — Candidatura duplicada: o sistema informa o registro existente.", "A3 — Perfil incompleto: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA] se deve bloquear."],
        "Tela de detalhe da vaga (/jobs/[id])"),
    make_uc("UC023", "Participar de evento", "Condução", "Usuário / Estudante / Desenvolvedor", "Stacklyst",
        "Inscrever usuário em evento comunitário elegível.", "Usuário autenticado; evento ativo para inscrição.",
        "Participação registrada uma única vez.", "Identificador do evento.", "Confirmação, regras e estado da inscrição.",
        ["O usuário abre a listagem e seleciona um evento.", "O Stacklyst apresenta período, nível e capacidade.", "O usuário solicita participação.", "O sistema valida período, nível, limite e inscrição prévia.", "O Stacklyst registra o participante e confirma a inscrição."],
        ["A1 — Nível insuficiente: o sistema informa o requisito.", "A2 — Evento lotado ou encerrado: o sistema recusa a inscrição.", "A3 — Inscrição duplicada: o sistema mantém o registro existente."],
        "Tela de eventos (/events)"),
    make_uc("UC024", "Solicitar função de avaliador", "Condução", "Usuário / Estudante / Desenvolvedor", "Administrador; Stacklyst",
        "Verificar elegibilidade e enviar candidatura para papel de avaliador.", "Usuário autenticado e sem candidatura pendente.",
        "Candidatura pendente criada ou inelegibilidade informada.", "Motivação e tecnologias de domínio.", "Elegibilidade e status da candidatura.",
        ["O usuário abre a área de candidatura a avaliador.", "O Stacklyst calcula trilhas completas e XP.", "Se elegível, o usuário informa motivação e tecnologias.", "O sistema valida e registra a candidatura pendente.", "O administrador passa a visualizar a candidatura para decisão."],
        ["A1 — Critério não atendido: o sistema mostra requisitos atuais.", "A2 — Candidatura já pendente: o sistema não duplica.", "A3 — Usuário já é avaliador/administrador: o sistema informa o papel existente."],
        "Tela de candidatura a avaliador (/evaluators/apply)"),
    make_uc("UC025", "Cadastrar empresa", "Configuração", "Empresa / Recrutador; Administrador", "Stacklyst",
        "Criar perfil de empresa vinculado a responsável autorizado.", "Ator autenticado com papel RECRUITER ou ADMIN.",
        "Empresa criada e disponível para gestão de vagas/eventos.", "Nome, slug, descrição, localização, site e imagens opcionais.", "Cadastro da empresa.",
        ["O ator abre a área de cadastro de empresa.", "O Stacklyst verifica o papel do solicitante.", "O ator informa os dados institucionais disponíveis.", "O sistema valida unicidade, formato e vínculo do proprietário.", "O Stacklyst cria a empresa sem marcar verificação oficial automaticamente."],
        ["A1 — Sem permissão: o sistema nega o cadastro.", "A2 — Slug/nome conflitante: o sistema solicita alteração.", "A3 — Verificação oficial: [ANOTAÇÃO — REGRA PRECISA SER DEFINIDA] para documentos e aprovação."],
        "Tela de cadastro/gestão de empresa (/recruiter)"),
]


TRACEABILITY = [
    ["AN01 — Realizar aprendizado por trilha", "RF005–RF010", "RN001–RN003", "UC004–UC007"],
    ["AN02 — Participar de duelo", "RF016–RF021", "RN007–RN013, RN028", "UC008–UC010"],
    ["AN03 — Resolver exercício", "RF008–RF009", "RN001–RN002, RN012", "UC006"],
    ["AN04 — Evoluir na gamificação", "RF011–RF015", "RN001–RN006, RN028", "UC007, UC011"],
    ["AN05 — Interagir com a comunidade", "RF030–RF033", "RN018–RN019, RN022", "UC013–UC014, UC020"],
    ["AN06 — Utilizar assistência da IA", "RF022–RF025", "RN013, RN024, RN027", "UC012"],
    ["AN07 — Atuar como avaliador", "RF026–RF029", "RN012–RN017", "UC010, UC016, UC024"],
    ["AN08 — Empresa divulgar vaga", "RF034–RF038, RF047–RF048", "RN020–RN021, RN026–RN027", "UC017–UC018, UC022, UC025"],
    ["AN09 — Administração da plataforma", "RF039–RF045, RF049", "RN015, RN017, RN019, RN022–RN025", "UC015, UC019–UC021, UC023–UC024"],
]


def prototype_for(document: Document, prototype: str, use_case_id: str) -> None:
    current_prototype = CURRENT_PROTOTYPES / prototype
    legacy_prototype = SCREENSHOTS / prototype
    if prototype.endswith(".png") and current_prototype.exists():
        caption = f"Protótipo de interface relacionado ao {use_case_id} — captura atual da plataforma autenticada em 22/08/2026."
        add_figure(document, current_prototype, caption, 6.2, 4.4)
    elif prototype.endswith(".png") and legacy_prototype.exists():
        add_figure(document, legacy_prototype, f"Protótipo de interface relacionado ao {use_case_id} — captura existente no repositório.", 6.2, 4.4)
    else:
        add_note(document, "[INSERIR PROTÓTIPO DA TELA CORRESPONDENTE]", prototype, "blue", spacing=False)


def build_use_cases() -> None:
    document = new_from_template("use_cases", "Especificação de Casos de Uso", "Especificação de Casos de Uso do Sistema")
    document.add_heading("2. Objetivo", level=1)
    document.add_paragraph(
        "Este documento detalha as interações entre os atores e as funcionalidades do Stacklyst. Os casos de uso correspondem ao diagrama PlantUML, aos requisitos funcionais e às atividades de negócio, incluindo fluxos alternativos e evidências de interface."
    )
    document.add_heading("3. Atores do Sistema", level=1)
    actor_rows = [
        ["Usuário / Estudante / Desenvolvedor", "Ator humano", "Usa trilhas, lições, exercícios, desafios, comunidade, gamificação, eventos, vagas e notificações próprias."],
        ["Avaliador", "Ator humano especializado", "Analisa soluções que exigem avaliação humana após elegibilidade e aprovação administrativa."],
        ["Administrador", "Ator humano privilegiado", "Gerencia usuários, papéis, denúncias, avaliadores, eventos e configurações suportadas."],
        ["Empresa / Recrutador", "Ator humano/organizacional", "Cadastra empresa, publica vagas/eventos e acessa somente informações permitidas."],
        ["Serviço de IA", "Sistema externo/secundário", "Auxilia em orientação, geração de conteúdo e análise; não é autoridade final."],
        ["Executor de código", "Sistema externo/secundário", "Executa código com limites e retorna saída/testes; sua indisponibilidade deve ser tratada."],
    ]
    add_table(document, ["Ator", "Natureza", "Permissões e responsabilidades"], actor_rows, [2.0, 1.25, 3.5], 8)

    document.add_heading("4. Diagrama de Casos de Uso", level=1)
    add_figure(document, DIAGRAMS / "imagens" / "casos-de-uso-stacklyst.png", "Figura 1 — Diagrama completo de casos de uso do Stacklyst.", 6.55, 8.2)
    document.add_page_break()
    document.add_heading("4.1 Código PlantUML", level=2)
    add_code_block(document, (DIAGRAMS / "fontes" / "casos-de-uso-stacklyst.puml").read_text(encoding="utf-8"))

    document.add_heading("5. Especificação dos Casos de Uso", level=1)
    for index, use_case in enumerate(USE_CASES, 1):
        document.add_heading(f"5.{index} Caso de Uso {use_case['id']} — {use_case['name']}", level=2)
        add_labeled(document, "Descrição", str(use_case["description"]))
        add_labeled(document, "Tipo", str(use_case["type"]))
        add_labeled(document, "Atores que iniciam", str(use_case["initiators"]))
        add_labeled(document, "Atores secundários", str(use_case["secondary"]))
        add_labeled(document, "Pré-condições", str(use_case["preconditions"]))
        add_labeled(document, "Pós-condições", str(use_case["postconditions"]))
        add_labeled(document, "Entradas", str(use_case["inputs"]))
        add_labeled(document, "Saídas", str(use_case["outputs"]))
        document.add_heading("Fluxo Principal", level=3)
        add_numbered(document, use_case["flow"])
        document.add_heading("Fluxos Alternativos", level=3)
        for alternative in use_case["alternatives"]:
            add_bullets(document, [str(alternative)])
        document.add_heading("Protótipo de Interface", level=3)
        prototype_for(document, str(use_case["prototype"]), str(use_case["id"]))
        if index != len(USE_CASES):
            document.add_page_break()

    document.add_heading("6. Anotações e Pendências do Stacklyst", level=1)
    document.add_heading("6.1 🔴 Decisões obrigatórias", level=2)
    for text in [
        "Definir o recorte oficial do MVP e separar mobile, recomendação por IA e recrutamento analítico em fases.",
        "Resolver a divergência do matchmaking: código atual amplia para ±1.000 XP e depois qualquer usuário; a proposta inicial menciona faixa superior.",
        "Definir resultado server-side quando convite ou desafio expirar e quando somente um participante enviar solução.",
        "Definir fórmula de efeito do duelo no ranking, empate, abandono, fraude, anulação e recurso.",
        "Definir campos públicos/privados, consentimento e auditoria do acesso de empresas aos indicadores.",
        "Definir critérios de publicação e revisão de exercícios/recomendações gerados por IA.",
        "Definir política de moderação, sanções, prazo de análise e recurso do usuário.",
        "Preencher integrantes, datas, orçamento e confirmar a caracterização acadêmica sem cliente real.",
    ]:
        add_note(document, "🔴 [DECISÃO NECESSÁRIA]", text, "red")
    document.add_heading("6.2 🟡 Itens para validação", level=2)
    for text in [
        "Validar baselines atuais: convite de 30 s, 3 rejeições, cooldown de 5 min e duelo de 15 min.",
        "Validar faixas de XP, divisões, sequência e critérios de conquistas.",
        "Validar elegibilidade de avaliador: uma trilha completa ou 1.000 XP, seguida de aprovação administrativa.",
        "Confirmar técnicas de elicitação realmente utilizadas e os registros que comprovam sua realização.",
        "Definir métricas de desempenho, disponibilidade, acessibilidade, carga e navegadores suportados.",
        "Validar ações administrativas completas de conteúdo, eventos, auditoria e notificações aos envolvidos.",
        "Revisar com a professora se Mermaid é aceito como representação dos diagramas de atividades.",
    ]:
        add_note(document, "🟡 [VALIDAR]", text, "yellow")
    document.add_heading("6.3 🔵 Sugestões futuras", level=2)
    for text in [
        "Entregar em fases: núcleo educacional; comunidade/gamificação; duelos/avaliação; vagas; mobile.",
        "Usar feature flags e fallbacks para que IA e serviços externos permaneçam opcionais.",
        "Criar tabela específica de auditoria e política de retenção para ações sensíveis.",
        "Publicar contratos de API e testes de contrato antes do aplicativo mobile.",
        "Executar testes de carga representativos para feed, ranking, matchmaking e notificações.",
        "Realizar avaliação de acessibilidade e usabilidade com estudantes/desenvolvedores.",
    ]:
        add_note(document, "🔵 [SUGESTÃO]", text, "blue")

    document.add_heading("7. Matriz de Rastreabilidade Geral", level=1)
    add_table(document, ["Atividade do Negócio", "Requisito Funcional", "Regra de Negócio", "Caso de Uso"], TRACEABILITY, [2.2, 1.65, 1.65, 1.25], 7.3)

    document.add_heading("8. Revisão de Consistência com a Programação", level=1)
    implementation_rows = [
        ["Cadastro/login/perfil", "Implementado", "src/app/register; src/app/login; src/app/profile; src/app/api/auth/profile"],
        ["Recuperação de acesso", "Planejado", "Nenhum fluxo completo localizado"],
        ["Trilhas/lições/exercícios", "Implementado", "src/app/trails; src/app/lesson; src/lib/lessons"],
        ["XP/níveis/conquistas", "Implementado com regras a validar", "src/services/xp.service.ts; src/lib/streak.ts"],
        ["Ranking/divisões", "Parcial", "leaderboard implementado; divisões usadas no duelo; integração competitiva incompleta"],
        ["Duelos/matchmaking", "Parcial", "convite e busca existem; fila temporal, expiração server-side e efeito no ranking pendentes"],
        ["Avaliação humana", "Implementado", "evaluator applications, RBAC e DuelEvaluation"],
        ["IA educacional", "Parcial e opcional", "chat/análise existem; personalização persistente e recomendações não"],
        ["Comunidade", "Implementado", "feed, posts, respostas, reações, votos, favoritos, denúncias e mensagens"],
        ["Eventos", "Parcial", "criação e participação existem; ciclo administrativo completo precisa de validação"],
        ["Vagas/empresas", "Parcial", "cadastro, vaga, candidatura e etapas existem; política de acesso a indicadores pendente"],
        ["Administração", "Parcial", "usuários, denúncias, avaliadores e métricas existem; auditoria abrangente pendente"],
        ["Notificações", "Implementado na web", "central e tipos de notificação; push mobile ausente"],
        ["Aplicação mobile", "Planejado", "nenhum projeto mobile localizado no repositório"],
    ]
    add_table(document, ["Área", "Estado", "Evidência/limite"], implementation_rows, [1.75, 1.45, 3.55], 7.5)
    document.add_paragraph(
        "Conclusão da revisão: os quatro documentos usam os mesmos atores, identificadores e regras. Diferenças entre a proposta e o código foram tratadas como decisões/validações, sem transformar funcionalidade planejada em implementação concluída."
    )
    document.save(FINAL_FILES["use_cases"])


def write_manifest() -> None:
    originals = [
        Path(r"C:\Users\PEDRO\Downloads\1. Documento Visão (1).doc"),
        Path(r"C:\Users\PEDRO\Downloads\2. Atividades do Negócio (1).doc"),
        Path(r"C:\Users\PEDRO\Downloads\3. Requisitos do Sistema (1).doc"),
        Path(r"C:\Users\PEDRO\Downloads\4. Casos de Uso (1).doc"),
    ]
    payload = {
        "project": "Stacklyst",
        "date": "[ANOTAÇÃO — DATA A DEFINIR]",
        "original_templates": [
            {"path": str(path), "sha256": sha256(path), "preserved": True} for path in originals
        ],
        "deliverables": [
            {"path": str(path.relative_to(REPO)), "sha256": sha256(path)} for path in FINAL_FILES.values()
        ],
        "figjam": "https://www.figma.com/board/kCKbpkIxJiQHmcprLH3Tfd",
    }
    (ROOT / "manifesto-de-integridade.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    for key, path in TEMPLATES.items():
        if not path.exists():
            raise FileNotFoundError(f"Cópia de trabalho do template {key} não encontrada: {path}")
    build_vision()
    build_activities()
    build_requirements()
    build_use_cases()
    write_manifest()
    for path in FINAL_FILES.values():
        print(f"GENERATED {path} {path.stat().st_size} bytes SHA256={sha256(path)}")


if __name__ == "__main__":
    main()
